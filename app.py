"""
GP_Partner_Ultimate - 高级投资辅助系统
支持多模式切换、OpenRouter API、历史记录和报告导出
"""

import streamlit as st
import os
import json
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, List, Union
import re
import zipfile
import io
import hashlib
import time

# 第三方库导入
try:
    from openai import OpenAI
    from dotenv import load_dotenv
    import PyPDF2
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import pandas as pd
except ImportError as e:
    st.error(f"❌ 缺少必要的依赖库: {e}")
    st.stop()

# 注：已放弃向量检索方案，改用基于 LLM 标签提取的轻量级记忆系统

# 加载环境变量
load_dotenv()

# 页面配置
st.set_page_config(
    page_title="GP Partner Ultimate",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 常量定义
PROMPTS_DIR = Path("prompts")
HISTORY_DIR = Path("history_data")
BASE_URL = "https://openrouter.ai/api/v1"
DEFAULT_MODEL = "google/gemini-flash-1.5"
KNOWLEDGE_BASE_FILE = Path("project_database.csv")
MEMORY_STORE_FILE = Path("memory_store.json")  # 基于标签的记忆存储文件
EVOLUTION_LOG_FILE = Path("evolution_log.md")  # AI 进化日志文件
CONFIG_FILE = Path("config.json")  # 配置持久化文件

# 确保必要的文件夹存在
HISTORY_DIR.mkdir(exist_ok=True)
PROMPTS_DIR.mkdir(exist_ok=True)


# ==================== 配置持久化系统 ====================

def load_config() -> Dict:
    """加载配置文件"""
    try:
        if CONFIG_FILE.exists():
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                config = json.load(f)
                return config
        else:
            # 返回默认配置
            return {
                "api_key": os.getenv("OPENROUTER_API_KEY", ""),
                "base_url": BASE_URL,
                "model": DEFAULT_MODEL
            }
    except Exception as e:
        print(f"⚠️ 加载配置文件失败: {e}")
        return {
            "api_key": os.getenv("OPENROUTER_API_KEY", ""),
            "base_url": BASE_URL,
            "model": DEFAULT_MODEL
        }


def save_config():
    """保存配置到文件"""
    try:
        config = {
            "api_key": st.session_state.get("api_key_input", ""),
            "base_url": st.session_state.get("base_url_input", BASE_URL),
            "model": st.session_state.get("model_input", DEFAULT_MODEL)
        }
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"⚠️ 保存配置文件失败: {e}")


# ==================== 基于 LLM 标签提取的轻量级记忆系统 ====================

def extract_tags_from_text(text: str, client: OpenAI, model: str) -> List[str]:
    """使用 LLM 从文本中提取项目标签"""
    try:
        tag_extraction_prompt = f"""请从以下项目文本中提取3-5个核心标签（关键词），用于项目分类和匹配。

文本内容：
{text[:1000]}

要求：
1. 标签应该是项目的核心技术、行业、商业模式等关键特征
2. 返回格式：纯文本，用逗号分隔，不要使用任何标记符号
3. 例如：AI, 医疗, B2B, SaaS, 多模态

标签："""

        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "你是一个专业的项目标签提取助手。只需返回标签，用逗号分隔，不要添加任何解释。"},
                {"role": "user", "content": tag_extraction_prompt}
            ],
            temperature=0.3,
            max_tokens=100
        )
        
        tags_text = response.choices[0].message.content.strip()
        # 解析标签
        tags = [tag.strip() for tag in tags_text.split(',') if tag.strip()]
        return tags[:5]  # 最多返回5个标签
    except Exception as e:
        print(f"⚠️ 标签提取失败: {e}")
        return []


class MemoryManager:
    """基于 LLM 标签提取的轻量级记忆系统管理器（单例模式）"""
    _instance = None
    _initialized = False
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super(MemoryManager, cls).__new__(cls)
        return cls._instance
    
    def __init__(self):
        if self._initialized:
            return
        
        self.memory_store_path = MEMORY_STORE_FILE
        self.memories = []
        self.enabled = True
        
        # 加载记忆库
        self._load_memories()
        self._initialized = True
    
    def _load_memories(self):
        """从 JSON 文件加载记忆库"""
        try:
            if self.memory_store_path.exists():
                with open(self.memory_store_path, 'r', encoding='utf-8') as f:
                    self.memories = json.load(f)
            else:
                self.memories = []
        except Exception as e:
            print(f"⚠️ 加载记忆库失败: {e}")
            self.memories = []
    
    def _save_memories(self):
        """保存记忆库到 JSON 文件"""
        try:
            with open(self.memory_store_path, 'w', encoding='utf-8') as f:
                json.dump(self.memories, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            print(f"⚠️ 保存记忆库失败: {e}")
            return False
    
    def query_similar(self, text: str, client: OpenAI, model: str, top_k: int = 3) -> List[Dict]:
        """基于标签匹配查询相似的历史项目"""
        if not self.enabled or len(self.memories) == 0:
            return []
        
        try:
            # 提取当前项目的标签
            query_tags = extract_tags_from_text(text, client, model)
            if not query_tags:
                return []
            
            # 计算每个历史项目的标签匹配度
            scored_projects = []
            for memory in self.memories:
                memory_tags = memory.get('tags', [])
                if not memory_tags:
                    continue
                
                # 计算标签交集数量（简单的标签匹配）
                common_tags = set(query_tags) & set(memory_tags)
                match_score = len(common_tags) / max(len(query_tags), len(memory_tags)) if max(len(query_tags), len(memory_tags)) > 0 else 0
                
                if match_score > 0:  # 至少有一个共同标签
                    scored_projects.append({
                        'name': memory.get('id', '未知项目'),
                        'score': memory.get('score', 'N/A'),
                        'summary': memory.get('summary', ''),
                        'match_score': match_score,
                        'common_tags': list(common_tags)
                    })
            
            # 按匹配度排序，返回前 top_k 个
            scored_projects.sort(key=lambda x: x['match_score'], reverse=True)
            return scored_projects[:top_k]
        except Exception as e:
            print(f"⚠️ 查询相似项目失败: {e}")
            return []
    
    def add_memory(self, name: str, summary: str, full_text: str, score, tags: List[str] = None, meta: Dict = None):
        """添加项目到记忆库"""
        if not self.enabled:
            return False
        
        try:
            # 如果未提供标签，使用空列表（标签会在外部通过 LLM 提取）
            if tags is None:
                tags = []
            
            memory_entry = {
                "id": name,
                "summary": summary,
                "score": score,
                "tags": tags,
                "timestamp": datetime.now().isoformat()
            }
            
            if meta:
                memory_entry.update(meta)
            
            # 检查是否已存在同名项目，如果存在则更新，否则添加
            existing_index = None
            for i, mem in enumerate(self.memories):
                if mem.get('id') == name:
                    existing_index = i
                    break
            
            if existing_index is not None:
                self.memories[existing_index] = memory_entry
            else:
                self.memories.append(memory_entry)
            
            # 保存到文件
            return self._save_memories()
        except Exception as e:
            print(f"⚠️ 添加记忆失败: {e}")
            return False
    
    def get_count(self) -> int:
        """获取知识库中的项目数量"""
        return len(self.memories) if self.enabled else 0


# 全局 MemoryManager 实例
memory_manager = MemoryManager()


def load_prompt_files() -> Dict[str, str]:
    """扫描 prompts 文件夹，加载所有 .txt 文件"""
    prompt_files = {}
    if not PROMPTS_DIR.exists():
        return prompt_files
    
    txt_files = list(PROMPTS_DIR.glob("*.txt"))
    for file_path in sorted(txt_files):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read().strip()
                # 使用文件名（不含扩展名）作为键
                name = file_path.stem
                prompt_files[name] = content
        except Exception as e:
            st.warning(f"⚠️ 无法读取文件 {file_path.name}: {e}")
    
    return prompt_files


def extract_text_from_pdf(file) -> str:
    """从 PDF 文件中提取文本"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"❌ PDF 解析错误: {e}")
        return ""


def extract_text_from_docx(file) -> str:
    """从 DOCX 文件中提取文本"""
    try:
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"❌ DOCX 解析错误: {e}")
        return ""


def extract_text_from_uploaded_file(uploaded_file) -> Optional[str]:
    """根据文件类型提取文本"""
    file_extension = Path(uploaded_file.name).suffix.lower()
    
    if file_extension == ".pdf":
        return extract_text_from_pdf(uploaded_file)
    elif file_extension in [".docx", ".doc"]:
        return extract_text_from_docx(uploaded_file)
    else:
        st.error(f"❌ 不支持的文件格式: {file_extension}")
        return None


def save_history_entry(mode_name: str, file_name: str, analysis_content: str):
    """保存历史记录到 JSON 文件"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    entry = {
        "mode": mode_name,
        "file_name": file_name,
        "timestamp": timestamp,
        "datetime": datetime.now().isoformat(),
        "content": analysis_content
    }
    
    filename = f"{timestamp}_{mode_name.replace(' ', '_')}.json"
    filepath = HISTORY_DIR / filename
    
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(entry, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        st.error(f"❌ 保存历史记录失败: {e}")
        return False


def load_history_entries() -> List[Dict]:
    """加载所有历史记录"""
    entries = []
    if not HISTORY_DIR.exists():
        return entries
    
    json_files = sorted(HISTORY_DIR.glob("*.json"), reverse=True)
    for filepath in json_files:
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                entry = json.load(f)
                entries.append(entry)
        except Exception as e:
            st.warning(f"⚠️ 无法读取历史记录 {filepath.name}: {e}")
    
    return entries


def clean_markdown_for_display(markdown_text: str) -> str:
    """清理 Markdown 文本，移除 Tags 头部、JSON 代码块和进化建议（用于显示和导出）"""
    cleaned = markdown_text
    # 移除 Tags 头部
    cleaned = re.sub(r'---TAGS:\s*\[[^\]]+\]---\s*\n?', '', cleaned)
    # 移除 ```json ... ``` 代码块
    pattern = r'```json\s*[\s\S]*?```'
    cleaned = re.sub(pattern, '', cleaned, flags=re.DOTALL)
    # 移除进化建议章节
    cleaned = re.sub(r'##\s*🧬\s*进化建议\s*\n.*?(?=\n```json|\n```|$)', '', cleaned, flags=re.DOTALL)
    # 清理多余的空行
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()


def markdown_to_docx(markdown_text: str, output_path: str):
    """将 Markdown 文本转换为 Word 文档"""
    # 清理 JSON 代码块
    markdown_text = clean_markdown_for_display(markdown_text)
    
    doc = Document()
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # 空行
            doc.add_paragraph()
            continue
        
        # 检查标题层级
        if line.startswith('# '):
            # H1 标题
            heading = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            # H2 标题
            heading = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            # H3 标题
            heading = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            # H4 标题
            heading = doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            para = doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.startswith(('  - ', '  * ', '    - ', '    * ')):
            # 嵌套列表项
            indent_level = len(line) - len(line.lstrip())
            para = doc.add_paragraph(line.lstrip()[2:].strip(), style='List Bullet')
            para.paragraph_format.left_indent = Pt(indent_level * 12)
        elif re.match(r'^\d+\.\s+', line):
            # 有序列表
            para = doc.add_paragraph(line, style='List Number')
        elif line.startswith('**') and line.endswith('**'):
            # 粗体文本
            para = doc.add_paragraph()
            para.add_run(line[2:-2]).bold = True
        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            # 斜体文本
            para = doc.add_paragraph()
            para.add_run(line[1:-1]).italic = True
        else:
            # 普通段落 - 处理行内格式（粗体、斜体等）
            para = doc.add_paragraph()
            # 处理粗体 **text**
            if '**' in line:
                parts = re.split(r'(\*\*[^\*]+?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        para.add_run(part[2:-2]).bold = True
                    elif part.strip():
                        # 处理斜体 *text*
                        if '*' in part and not part.startswith('*'):
                            italic_parts = re.split(r'(\*[^\*]+\*)', part)
                            for italic_part in italic_parts:
                                if italic_part.startswith('*') and italic_part.endswith('*') and len(italic_part) > 2:
                                    para.add_run(italic_part[1:-1]).italic = True
                                elif italic_part.strip():
                                    para.add_run(italic_part)
                        else:
                            para.add_run(part)
            else:
                # 没有粗体标记，直接添加文本
                para.add_run(line)
    
    try:
        doc.save(output_path)
        return True
    except Exception as e:
        st.error(f"❌ 保存 Word 文档失败: {e}")
        return False


def enhance_system_prompt(base_prompt: str, similar_projects: List[Dict] = None) -> str:
    """增强系统提示词，要求 LLM 严格按照特定格式输出"""
    
    # 输出格式要求
    format_instruction = """

---
【严格输出格式要求】

你的输出必须严格按照以下三部分结构：

**第一部分 - 标签头部（必须）**：
在内容最开始，输出一行：
---TAGS: ["标签1", "标签2", "标签3", "标签4", "标签5"]---

**第二部分 - 分析正文**：
正常的 Markdown 格式分析报告，包含所有分析内容。

**第三部分 - 进化建议（必须）**：
在报告最后，必须包含一个章节：
## 🧬 进化建议
[针对本次分析，反思当前 System Prompt 的不足，并给出具体的优化指令建议。建议应该具体、可操作，例如："建议在提示词中增加对XX赛道的特殊关注"或"建议明确要求分析XX维度的风险"等]

**第四部分 - JSON 数据块（必须）**：
在进化建议之后，附带一个 JSON 数据块：

```json
{
  "project_name": "项目名称",
  "industry": "所属赛道/行业（如：AI Agent, 具身智能）",
  "tags": ["技术标签1", "技术标签2", "技术标签3"],
  "stage": "融资阶段（如：Angel, Pre-A, A轮, B轮等）",
  "score": 8,
  "summary": "一句话核心评价（50字以内）",
  "risk_level": "High/Medium/Low"
}
```

注意：
- project_name: 从商业计划书中提取的项目名称（字符串）
- industry: 所属赛道/行业（字符串，如：AI Agent, 具身智能, SaaS, 区块链等）
- tags: 技术标签列表（数组，如：["RAG", "LLM", "SaaS", "多模态"]），必须与第一部分的 TAGS 保持一致
- stage: 融资阶段（字符串，如：Angel, Pre-A, A轮, B轮等，如果未提及则填写 "未披露"）
- score: 投资推荐评分（整数，1-10分，10分为最高），必须在正文中明确显示
- summary: 一句话核心评价（字符串，50字以内，简洁概括项目价值和风险）
- risk_level: 风险等级（字符串，必须是 "High", "Medium", "Low" 之一）

请严格按照以上格式输出，不要遗漏任何部分。
"""
    
    # 如果有相似项目，添加历史项目参考信息
    if similar_projects and len(similar_projects) > 0:
        similar_info = "\n\n---\n【历史项目参考】知识库中发现了与本项目相似的过往项目（基于标签匹配）：\n"
        for proj in similar_projects:
            common_tags_str = ', '.join(proj.get('common_tags', [])) if proj.get('common_tags') else '无'
            match_score = proj.get('match_score', 0)
            similar_info += f"- {proj['name']} (评分: {proj['score']}, 核心评价: {proj['summary']}, 共同标签: {common_tags_str}, 匹配度: {match_score:.2%})\n"
        similar_info += "\n请在分析时横向对比，指出本项目的差异化优势或重复造轮子的风险。参考这些历史案例，但不要被其局限，重点分析当前项目的独特价值。\n"
        return base_prompt + similar_info + format_instruction
    
    return base_prompt + format_instruction


def extract_tags_from_response(response_text: str) -> List[str]:
    """从 LLM 响应中提取 Tags（从头部 ---TAGS: ... --- 格式）"""
    try:
        # 匹配 ---TAGS: ["标签1", "标签2"]--- 格式
        pattern = r'---TAGS:\s*(\[[^\]]+\])---'
        match = re.search(pattern, response_text)
        
        if match:
            tags_str = match.group(1)
            # 解析 JSON 数组
            tags = json.loads(tags_str)
            if isinstance(tags, list):
                return [str(tag).strip() for tag in tags if tag]
        
        # 如果没有找到，尝试从 JSON 数据块中提取
        json_data = extract_json_from_response(response_text)
        if json_data and 'tags' in json_data:
            tags = json_data.get('tags', [])
            if isinstance(tags, list):
                return [str(tag).strip() for tag in tags if tag]
        
        return []
    except Exception as e:
        print(f"⚠️ 提取 Tags 失败: {e}")
        return []


def extract_evolution_suggestion(response_text: str) -> Optional[str]:
    """从 LLM 响应中提取进化建议（## 🧬 进化建议 后的内容）"""
    try:
        # 匹配 ## 🧬 进化建议 后的内容（直到 JSON 代码块或文件末尾）
        pattern = r'##\s*🧬\s*进化建议\s*\n(.*?)(?=\n```json|\n```|$)'
        match = re.search(pattern, response_text, re.DOTALL)
        
        if match:
            suggestion = match.group(1).strip()
            # 清理多余的空行
            suggestion = re.sub(r'\n{3,}', '\n\n', suggestion)
            return suggestion if suggestion else None
        
        return None
    except Exception as e:
        print(f"⚠️ 提取进化建议失败: {e}")
        return None


def extract_score_enhanced(response_text: str, json_data: Optional[Dict] = None) -> Optional[Union[int, float, str]]:
    """增强的分数提取函数，支持多种格式"""
    # 首先尝试从 JSON 数据中提取
    if json_data and 'score' in json_data:
        score = json_data.get('score')
        if score and score != 'N/A':
            try:
                if isinstance(score, (int, float)):
                    return score
                score_str = str(score).strip()
                # 尝试转换为数字
                if score_str.replace('.', '').isdigit():
                    return float(score_str) if '.' in score_str else int(score_str)
            except:
                pass
    
    # 在前 1000 字符中搜索分数
    search_text = response_text[:1000]
    
    # 模式1: Score: 8 或 评分：8.5
    patterns = [
        r'(?:Score|评分|分数|投资评分)[:：]\s*(\d+(?:\.\d+)?)',
        r'\[(\d+(?:\.\d+)?)分\]',
        r'评分[：:]\s*(\d+(?:\.\d+)?)',
        r'(\d+(?:\.\d+)?)\s*分',
        r'(\d+(?:\.\d+)?)\s*/\s*10',  # x/10 格式
    ]
    
    for pattern in patterns:
        match = re.search(pattern, search_text, re.IGNORECASE)
        if match:
            try:
                score_val = float(match.group(1))
                if 1 <= score_val <= 10:
                    return int(score_val) if score_val.is_integer() else score_val
            except:
                continue
    
    # 在整个文本中搜索 x/10 格式（兜底策略）
    x_10_pattern = r'(\d+(?:\.\d+)?)\s*/\s*10'
    matches = re.findall(x_10_pattern, response_text)
    if matches:
        for match in matches:
            try:
                score_val = float(match)
                if 1 <= score_val <= 10:
                    return int(score_val) if score_val.is_integer() else score_val
            except:
                continue
    
    return None


def parse_llm_response(response_text: str) -> Dict:
    """统一解析 LLM 响应，提取所有结构化信息"""
    result = {
        'tags': [],
        'json_data': None,
        'score': None,
        'evolution_suggestion': None,
        'body_content': response_text  # 原始内容
    }
    
    # 提取 Tags
    result['tags'] = extract_tags_from_response(response_text)
    
    # 提取 JSON 数据
    result['json_data'] = extract_json_from_response(response_text)
    
    # 提取分数（增强版）
    result['score'] = extract_score_enhanced(response_text, result['json_data'])
    
    # 提取进化建议
    result['evolution_suggestion'] = extract_evolution_suggestion(response_text)
    
    # 提取正文内容（移除 Tags 头部、进化建议和 JSON）
    body = response_text
    # 移除 Tags 头部
    body = re.sub(r'---TAGS:\s*\[[^\]]+\]---\s*\n?', '', body)
    # 移除进化建议章节
    body = re.sub(r'##\s*🧬\s*进化建议\s*\n.*?(?=\n```json|\n```|$)', '', body, flags=re.DOTALL)
    # 移除 JSON 代码块
    body = re.sub(r'```json\s*[\s\S]*?```', '', body)
    result['body_content'] = body.strip()
    
    return result


def save_evolution_suggestion(project_name: str, suggestion: str):
    """保存进化建议到日志文件"""
    try:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_entry = f"\n\n---\n## {timestamp} - {project_name}\n\n{suggestion}\n"
        
        with open(EVOLUTION_LOG_FILE, 'a', encoding='utf-8') as f:
            f.write(log_entry)
        return True
    except Exception as e:
        print(f"⚠️ 保存进化建议失败: {e}")
        return False


def extract_json_from_response(response_text: str) -> Optional[Dict]:
    """从 AI 响应中提取 JSON 数据块"""
    try:
        # 首先尝试匹配 ```json ... ``` 代码块（支持多行）
        pattern = r'```json\s*([\s\S]*?)\s*```'
        matches = re.findall(pattern, response_text)
        
        if matches:
            # 取最后一个匹配（通常 JSON 在末尾）
            json_str = matches[-1].strip()
            try:
                # 解析 JSON
                json_data = json.loads(json_str)
                # 验证是否包含必需字段
                required_fields = ['project_name', 'industry', 'tags', 'stage', 'score', 'summary', 'risk_level']
                if all(field in json_data for field in required_fields):
                    return json_data
            except json.JSONDecodeError:
                # JSON 解析失败，尝试清理后再解析
                # 移除可能的注释或额外字符
                json_str = re.sub(r'//.*', '', json_str)  # 移除单行注释
                json_str = re.sub(r'/\*.*?\*/', '', json_str, flags=re.DOTALL)  # 移除多行注释
                try:
                    json_data = json.loads(json_str)
                    required_fields = ['project_name', 'industry', 'tags', 'stage', 'score', 'summary', 'risk_level']
                    if all(field in json_data for field in required_fields):
                        return json_data
                except:
                    pass
        
        # 如果没有找到代码块，尝试直接查找 JSON 对象（兜底方案）
        # 使用更精确的正则表达式匹配完整的 JSON 对象
        json_pattern = r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}'
        json_matches = re.findall(json_pattern, response_text, re.DOTALL)
        
        if json_matches:
            # 从后往前尝试解析（通常 JSON 在末尾）
            for json_str in reversed(json_matches):
                try:
                    json_data = json.loads(json_str)
                    # 验证是否包含必需字段
                    required_fields = ['project_name', 'industry', 'tags', 'stage', 'score', 'summary', 'risk_level']
                    if all(field in json_data for field in required_fields):
                        return json_data
                except:
                    continue
        
        return None
    except Exception as e:
        # 静默失败，不中断主流程
        print(f"⚠️ 提取 JSON 失败: {e}")
        return None


def save_to_knowledge_base(json_data: Dict):
    """将提取的 JSON 数据保存到项目知识库 CSV 文件"""
    try:
        # 添加时间戳
        json_data['timestamp'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 将 tags 列表转换为字符串（CSV 格式，用分号分隔）
        if isinstance(json_data.get('tags'), list):
            json_data['tags'] = ';'.join(json_data['tags'])
        
        # 定义 CSV 列顺序
        columns = ['timestamp', 'project_name', 'industry', 'tags', 'stage', 'score', 'summary', 'risk_level']
        
        # 确保所有字段都存在
        for col in columns:
            if col not in json_data:
                json_data[col] = ''
        
        # 创建 DataFrame
        df_new = pd.DataFrame([json_data])
        
        # 检查文件是否存在
        if KNOWLEDGE_BASE_FILE.exists():
            # 读取现有数据
            df_existing = pd.read_csv(KNOWLEDGE_BASE_FILE, encoding='utf-8-sig')
            # 合并数据
            df_combined = pd.concat([df_existing, df_new], ignore_index=True)
        else:
            # 创建新文件
            df_combined = df_new
        
        # 按照列顺序重新排列
        df_combined = df_combined[columns]
        
        # 保存到 CSV（使用 utf-8-sig 编码以支持中文，避免 Windows Excel 乱码）
        df_combined.to_csv(KNOWLEDGE_BASE_FILE, index=False, encoding='utf-8-sig')
        return True
    except Exception as e:
        # 静默失败，不中断主流程
        print(f"⚠️ 保存知识库失败: {e}")
        return False


def load_knowledge_base() -> pd.DataFrame:
    """加载项目知识库"""
    try:
        if KNOWLEDGE_BASE_FILE.exists():
            df = pd.read_csv(KNOWLEDGE_BASE_FILE, encoding='utf-8-sig')
            # 确保 timestamp 列为字符串类型，便于格式化
            if 'timestamp' in df.columns:
                df['timestamp'] = df['timestamp'].astype(str)
            return df
        else:
            return pd.DataFrame()
    except Exception as e:
        print(f"⚠️ 加载知识库失败: {e}")
        return pd.DataFrame()


def format_timestamp_for_display(timestamp_str: str) -> str:
    """将时间戳格式化为 YYYY-MM-DD 格式"""
    try:
        # 尝试解析多种时间格式
        if isinstance(timestamp_str, str):
            # 处理 "YYYY-MM-DD HH:MM:SS" 格式
            if ' ' in timestamp_str:
                date_part = timestamp_str.split(' ')[0]
                return date_part
            # 处理 ISO 格式
            elif 'T' in timestamp_str:
                date_part = timestamp_str.split('T')[0]
                return date_part
            # 如果已经是 YYYY-MM-DD 格式
            elif len(timestamp_str) >= 10:
                return timestamp_str[:10]
        return timestamp_str
    except:
        return timestamp_str


def get_recent_projects(limit: int = 5) -> List[Dict]:
    """获取最近的项目记录"""
    try:
        df = load_knowledge_base()
        if df.empty:
            return []
        
        # 按时间倒序排列，取最新 N 条
        if 'timestamp' in df.columns and 'project_name' in df.columns:
            df_sorted = df.sort_values('timestamp', ascending=False)
            recent = df_sorted.head(limit)
            
            result = []
            for _, row in recent.iterrows():
                date_str = format_timestamp_for_display(str(row.get('timestamp', '')))
                project_name = str(row.get('project_name', '未知项目'))
                result.append({
                    'date': date_str,
                    'name': project_name
                })
            return result
        return []
    except Exception as e:
        print(f"⚠️ 获取最近项目失败: {e}")
        return []


def calculate_kb_statistics(df: pd.DataFrame) -> Dict:
    """计算知识库统计指标"""
    stats = {
        'total_projects': 0,
        'avg_score': 0.0,
        'top_industry': 'N/A',
        'industry_count': {}
    }
    
    try:
        if df.empty:
            return stats
        
        stats['total_projects'] = len(df)
        
        # 计算平均评分
        if 'score' in df.columns:
            try:
                scores = pd.to_numeric(df['score'], errors='coerce').dropna()
                if len(scores) > 0:
                    stats['avg_score'] = round(scores.mean(), 1)
            except:
                pass
        
        # 计算最热赛道
        if 'industry' in df.columns:
            industry_counts = df['industry'].value_counts()
            if len(industry_counts) > 0:
                stats['top_industry'] = industry_counts.index[0]
                stats['industry_count'] = industry_counts.to_dict()
        
        return stats
    except Exception as e:
        print(f"⚠️ 计算统计指标失败: {e}")
        return stats


def generate_file_id(file_name: str) -> str:
    """为文件生成唯一 ID"""
    # 使用文件名和当前时间戳生成唯一 ID
    content = f"{file_name}_{time.time()}"
    return hashlib.md5(content.encode()).hexdigest()[:12]


def process_single_file(
    uploaded_file,
    file_id: str,
    system_prompt: str,
    api_key: str,
    model: str,
    selected_mode: str,
    row_container,
    similar_projects: List[Dict] = None
) -> Optional[Dict]:
    """处理单个文件，返回结果字典"""
    try:
        # 提取文件文本
        file_text = extract_text_from_uploaded_file(uploaded_file)
        if not file_text:
            with row_container.container():
                st.error(f"❌ 无法从文件 {uploaded_file.name} 中提取文本")
            return None
        
        # 初始化 OpenAI 客户端（使用侧边栏配置的值）
        # 注意：base_url 和 api_key 应该从侧边栏的输入框获取
        # 这里需要从调用函数的参数中获取，或从 session_state 读取
        client = OpenAI(
            base_url=BASE_URL,  # Base URL 是固定的，使用常量
            api_key=api_key
        )
        
        # system_prompt 已经在外部增强（包含 RAG 信息），直接使用
        enhanced_prompt = system_prompt
        
        # 调用 API（流式）
        stream = call_openrouter_api(client, enhanced_prompt, file_text, model)
        if not stream:
            with row_container.container():
                st.error(f"❌ API 调用失败: {uploaded_file.name}")
            return None
        
        # 收集完整响应
        full_response = ""
        
        for chunk in stream:
            if chunk.choices[0].delta.content is not None:
                full_response += chunk.choices[0].delta.content
        
        # 解析 LLM 响应（提取 Tags、JSON、分数、进化建议等）
        parsed = parse_llm_response(full_response)
        json_data = parsed.get('json_data')
        extracted_tags = parsed.get('tags', [])
        extracted_score = parsed.get('score')
        evolution_suggestion = parsed.get('evolution_suggestion')
        body_content = parsed.get('body_content', full_response)
        
        # 如果 JSON 中有 tags，优先使用 JSON 中的（更准确）
        if json_data and json_data.get('tags'):
            extracted_tags = json_data.get('tags', [])
        
        # 如果 JSON 中有 score，优先使用 JSON 中的
        if json_data and json_data.get('score') and json_data.get('score') != 'N/A':
            extracted_score = json_data.get('score')
        
        # 保存到 CSV 知识库
        if json_data:
            save_to_knowledge_base(json_data)
        
        # 保存历史记录（保存完整内容）
        save_history_entry(selected_mode, uploaded_file.name, full_response)
        
        # 保存进化建议
        if evolution_suggestion:
            project_name = json_data.get('project_name', uploaded_file.name) if json_data else uploaded_file.name
            save_evolution_suggestion(project_name, evolution_suggestion)
        
        # 生成 Word 文档（内存中，使用清理后的正文内容）
        word_buffer = io.BytesIO()
        temp_path = HISTORY_DIR / f"temp_{file_id}.docx"
        # Word 导出时使用清理后的正文（不包含 Tags 头部和进化建议）
        markdown_to_docx(body_content, str(temp_path))
        
        with open(temp_path, "rb") as f:
            word_buffer.write(f.read())
        word_buffer.seek(0)
        
        # 清理临时文件
        try:
            temp_path.unlink()
        except:
            pass
        
        # 返回结果字典
        project_name = json_data.get('project_name', uploaded_file.name) if json_data else uploaded_file.name
        final_score = extracted_score if extracted_score is not None else (json_data.get('score', 'N/A') if json_data else 'N/A')
        
        return {
            'file_id': file_id,
            'file_name': uploaded_file.name,
            'markdown_content': body_content,  # 使用清理后的正文
            'full_response': full_response,  # 保留完整响应用于展示
            'word_buffer': word_buffer,
            'json_data': json_data,
            'project_name': project_name,
            'score': final_score,
            'tags': extracted_tags,
            'evolution_suggestion': evolution_suggestion,
            'similar_projects': similar_projects if similar_projects else [],
            'status': 'completed',
            'timestamp': datetime.now().isoformat()
        }
    
    except Exception as e:
        row_container.error(f"❌ 处理文件 {uploaded_file.name} 时出错: {str(e)}")
        return None


def render_result_row(result: Dict, index: int) -> bool:
    """渲染单个结果行，返回是否被选中"""
    col1, col2 = st.columns([0.05, 0.95])
    
    with col1:
        checkbox_key = f"select_{result['file_id']}"
        is_selected = st.checkbox(
            "",
            value=True,
            key=checkbox_key,
            label_visibility="collapsed"
        )
    
    with col2:
        # 构建 expander 标题
        score_display = f"{result['score']}分" if isinstance(result['score'], (int, float)) else str(result['score'])
        expander_title = f"✅ [{score_display}] {result['project_name']} (点击展开详情)"
        
        with st.expander(expander_title, expanded=False):
            # ========== 大脑思考路径展示 ==========
            st.info("🧠 **大脑思考路径**")
            
            # 显示提取到的标签
            tags = result.get('tags', [])
            if tags:
                tags_str = ', '.join([f"`{tag}`" for tag in tags])
                st.markdown(f"**核心关键词**: {tags_str}")
            else:
                st.markdown("**核心关键词**: 未提取到标签")
            
            # 显示记忆激活状态
            similar_projects = result.get('similar_projects', [])
            if similar_projects and len(similar_projects) > 0:
                project_names = [proj.get('name', '未知项目') for proj in similar_projects]
                common_tags_list = []
                for proj in similar_projects:
                    common_tags = proj.get('common_tags', [])
                    if common_tags:
                        common_tags_list.extend(common_tags)
                
                common_tags_str = ', '.join(set(common_tags_list)) if common_tags_list else '无'
                st.success(f"🎯 **记忆激活**: 发现与历史项目 `{', '.join(project_names)}` 存在关联（共同标签: {common_tags_str}），已进行横向对比。")
            else:
                st.info("🌱 **新物种收录**: 知识库中暂无同类，已作为种子存入大脑。")
            
            st.divider()
            
            # 显示清理后的 Markdown 内容（正文）
            cleaned_content = clean_markdown_for_display(result['markdown_content'])
            st.markdown(cleaned_content)
            
            # 单独下载按钮
            word_buffer = result.get('word_buffer')
            if word_buffer:
                word_buffer.seek(0)
                st.download_button(
                    label="📥 下载此报告",
                    data=word_buffer.read(),
                    file_name=f"{result['project_name']}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_row_{result['file_id']}"
                )
    
    return is_selected


def call_openrouter_api(client: OpenAI, system_prompt: str, user_content: str, model: str):
    """调用 OpenRouter API 进行流式响应"""
    try:
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content}
        ]
        
        stream = client.chat.completions.create(
            model=model,
            messages=messages,
            stream=True
        )
        
        return stream
    except Exception as e:
        st.error(f"❌ API 调用错误: {e}")
        return None


# ==================== 侧边栏 ====================
with st.sidebar:
    st.header("⚙️ 控制台")
    
    # 初始化 is_analyzing 状态
    if "is_analyzing" not in st.session_state:
        st.session_state.is_analyzing = False
    
    # 页面导航（如果正在分析则禁用）
    page = st.sidebar.radio(
        "页面导航",
        options=["🚀 AI 分析工作台", "📂 全量知识库", "📜 历史记录"],
        index=0,
        label_visibility="visible",
        disabled=st.session_state.is_analyzing
    )
    
    # 如果正在分析，显示警告
    if st.session_state.is_analyzing:
        st.warning("⚠️ 分析任务运行中，导航已锁定...")
    
    st.divider()
    
    # API 配置（持久化）
    st.subheader("API 配置")
    
    # 加载配置（仅在第一次时）
    if "config_loaded" not in st.session_state:
        config = load_config()
        st.session_state["api_key_input"] = config.get("api_key", os.getenv("OPENROUTER_API_KEY", ""))
        st.session_state["base_url_input"] = config.get("base_url", BASE_URL)
        st.session_state["model_input"] = config.get("model", DEFAULT_MODEL)
        st.session_state["config_loaded"] = True
        
        # 如果配置文件不存在，首次加载时保存默认配置
        if not CONFIG_FILE.exists():
            save_config()
    
    # Base URL（禁用，固定值，但保存在配置中）
    base_url = st.text_input(
        "Base URL",
        value=st.session_state.get("base_url_input", BASE_URL),
        disabled=True,
        key="base_url_input",
        help="OpenRouter API 基础 URL（固定）"
    )
    
    # API Key（支持自动保存）
    api_key = st.text_input(
        "API Key",
        value=st.session_state.get("api_key_input", ""),
        type="password",
        key="api_key_input",
        on_change=save_config,
        help="从配置文件加载或手动输入，修改后自动保存"
    )
    
    # Model（支持自动保存）
    model = st.text_input(
        "Model",
        value=st.session_state.get("model_input", DEFAULT_MODEL),
        key="model_input",
        on_change=save_config,
        help="OpenRouter 模型名称，修改后自动保存"
    )
    
    # 显示配置状态提示
    if CONFIG_FILE.exists():
        st.caption("💾 配置已自动保存")
    
    st.divider()
    
    # 分析模式选择（仅在工作台页面需要，且在未分析时显示）
    if page == "🚀 AI 分析工作台" and not st.session_state.is_analyzing:
        st.subheader("📋 分析模式")
        prompt_files = load_prompt_files()
        
        if not prompt_files:
            st.warning("⚠️ 请在 prompts 文件夹放入 .txt 提示词文件")
            selected_mode = None
            system_prompt = None
        else:
            mode_names = list(prompt_files.keys())
            selected_mode = st.selectbox(
                "选择分析模式",
                options=mode_names,
                index=0
            )
            base_prompt = prompt_files.get(selected_mode, "")
            # 存储基础提示词，稍后在调用 API 时增强
            system_prompt = base_prompt
    else:
        # 知识库页面不需要这些变量，但需要初始化以避免错误
        # 如果正在分析，从 session_state 恢复
        if st.session_state.is_analyzing and "selected_mode" in st.session_state:
            selected_mode = st.session_state.selected_mode
            system_prompt = st.session_state.get("system_prompt", "")
        else:
            selected_mode = None
            system_prompt = None
    
    st.divider()
    
    # AI 进化日志
    with st.expander("🧬 查看 AI 进化日志", expanded=False):
        try:
            if EVOLUTION_LOG_FILE.exists():
                with open(EVOLUTION_LOG_FILE, 'r', encoding='utf-8') as f:
                    log_content = f.read()
                if log_content.strip():
                    st.markdown(log_content)
                else:
                    st.info("暂无进化建议，分析项目后会自动生成。")
            else:
                st.info("暂无进化建议，分析项目后会自动生成。")
        except Exception as e:
            st.warning(f"⚠️ 读取进化日志失败: {e}")


# ==================== 主界面 ====================
# 根据页面导航显示不同内容
if page == "🚀 AI 分析工作台":
    # ========== AI 分析工作台页面 ==========
    st.title("🚀 AI 分析工作台")
    st.caption("高级投资辅助系统 - 支持批量分析与报告导出")

    # 显示当前模式
    if selected_mode:
        st.info(f"📋 **当前分析模式**: {selected_mode}")
    else:
        st.warning("⚠️ 请先在侧边栏选择分析模式")

    st.divider()

    # 初始化 session state
    if "processed_results" not in st.session_state:
        st.session_state["processed_results"] = []
    if "task_queue" not in st.session_state:
        st.session_state["task_queue"] = []
    if "processing_status" not in st.session_state:
        st.session_state["processing_status"] = {}

    # 初始化 MemoryManager（轻量级，无需额外依赖）
    # memory_manager 已在模块加载时自动初始化

    # 批量文件上传（分析中时禁用）
    uploaded_files = st.file_uploader(
        "📄 上传分析文件（支持批量）",
        type=["pdf", "docx", "doc"],
        help="支持 PDF 和 DOCX 格式，可同时上传多个文件进行批量分析",
        accept_multiple_files=True,
        disabled=st.session_state.is_analyzing
    )

    # 开始批量分析按钮（仅在未分析时显示）
    if uploaded_files and len(uploaded_files) > 0 and not st.session_state.is_analyzing:
        if st.button("🚀 开始批量分析", type="primary", disabled=not (selected_mode and api_key)):
            if not api_key:
                st.error("❌ 请输入 API Key")
                st.stop()
            
            if not selected_mode:
                st.error("❌ 请选择分析模式")
                st.stop()
            
            # 设置分析状态并保存配置到 session_state
            st.session_state.is_analyzing = True
            st.session_state.selected_mode = selected_mode
            st.session_state.system_prompt = system_prompt
            
            # 初始化任务队列
            task_queue = []
            for i, uploaded_file in enumerate(uploaded_files):
                file_id = generate_file_id(f"{uploaded_file.name}_{i}_{time.time()}")
                task_queue.append({
                    'file_id': file_id,
                    'file': uploaded_file,
                    'index': i
                })
            
            st.session_state.task_queue = task_queue
            st.rerun()
    
    # 如果正在分析，执行批量处理循环
    if st.session_state.is_analyzing and "task_queue" in st.session_state:
        task_queue = st.session_state.task_queue
        if len(task_queue) > 0:
            st.subheader("📋 任务队列")
            
            # 初始化所有行的占位符（仅在第一次渲染时）
            if "row_containers" not in st.session_state:
                st.session_state.row_containers = []
                for task in task_queue:
                    row_container = st.empty()
                    with row_container.container():
                        st.markdown(f"⏳ **等待处理**: {task['file'].name}")
                    st.session_state.row_containers.append({
                        'file_id': task['file_id'],
                        'container': row_container,
                        'status': 'waiting'
                    })
            
            # 获取当前处理索引
            if "current_task_index" not in st.session_state:
                st.session_state.current_task_index = 0
            
            current_idx = st.session_state.current_task_index
            
            if current_idx < len(task_queue):
                task = task_queue[current_idx]
                row_info = st.session_state.row_containers[current_idx]
                
                # 更新当前行的状态
                with row_info['container'].container():
                    st.markdown(f"🔄 **正在深入分析 (关联知识库中...)**: {task['file'].name}...")
                
                # 提取文件文本
                file_text = extract_text_from_uploaded_file(task['file'])
                
                if file_text:
                    # 初始化 OpenAI 客户端（用于标签提取和相似项目查询）
                    client = OpenAI(
                        base_url=base_url,  # 使用侧边栏配置的值
                        api_key=api_key
                    )
                    
                    # 基于标签匹配检索相似项目
                    similar_projects = []
                    if memory_manager.enabled and memory_manager.get_count() > 0:
                        similar_projects = memory_manager.query_similar(file_text, client, model, top_k=3)
                    
                    # 增强 Prompt（包含相似项目信息）
                    enhanced_prompt = enhance_system_prompt(
                        st.session_state.system_prompt,
                        similar_projects if similar_projects else None
                    )
                    
                    # 处理文件
                    result = process_single_file(
                        uploaded_file=task['file'],
                        file_id=task['file_id'],
                        system_prompt=enhanced_prompt,
                        api_key=api_key,
                        model=model,
                        selected_mode=st.session_state.selected_mode,
                        row_container=row_info['container'],
                        similar_projects=similar_projects
                    )
                    
                    if result:
                        # 提取 JSON 数据
                        json_data = result.get('json_data')
                        
                        # 存储到轻量级记忆库
                        if json_data and memory_manager.enabled:
                            project_name = json_data.get('project_name', task['file'].name)
                            summary = json_data.get('summary', '')
                            score = result.get('score', json_data.get('score', 'N/A'))
                            # 使用从响应中提取的 tags（已通过 parse_llm_response 提取）
                            tags = result.get('tags', [])
                            
                            # 如果仍然没有标签，尝试从 JSON 中获取
                            if not tags and json_data.get('tags'):
                                tags = json_data.get('tags', [])
                            
                            memory_manager.add_memory(
                                name=project_name,
                                summary=summary,
                                full_text=result.get('full_response', result['markdown_content']),  # 使用完整响应
                                score=score,
                                tags=tags,
                                meta={
                                    'industry': json_data.get('industry', ''),
                                    'stage': json_data.get('stage', ''),
                                    'risk_level': json_data.get('risk_level', '')
                                }
                            )
                        
                        # 更新行容器为最终结果
                        row_info['container'].empty()
                        with row_info['container'].container():
                            col1, col2 = st.columns([0.05, 0.95])
                            
                            with col1:
                                checkbox_key = f"select_{result['file_id']}"
                                st.checkbox(
                                    "",
                                    value=True,
                                    key=checkbox_key,
                                    label_visibility="collapsed"
                                )
                            
                            with col2:
                                score_display = f"{result['score']}分" if isinstance(result['score'], (int, float)) else str(result['score'])
                                expander_title = f"✅ [{score_display}] {result['project_name']} (点击展开详情)"
                                
                                with st.expander(expander_title, expanded=False):
                                    # ========== 大脑思考路径展示 ==========
                                    st.info("🧠 **大脑思考路径**")
                                    
                                    # 显示提取到的标签
                                    tags = result.get('tags', [])
                                    if tags:
                                        tags_str = ', '.join([f"`{tag}`" for tag in tags])
                                        st.markdown(f"**核心关键词**: {tags_str}")
                                    else:
                                        st.markdown("**核心关键词**: 未提取到标签")
                                    
                                    # 显示记忆激活状态
                                    similar_projects = result.get('similar_projects', [])
                                    if similar_projects and len(similar_projects) > 0:
                                        project_names = [proj.get('name', '未知项目') for proj in similar_projects]
                                        common_tags_list = []
                                        for proj in similar_projects:
                                            common_tags = proj.get('common_tags', [])
                                            if common_tags:
                                                common_tags_list.extend(common_tags)
                                        
                                        common_tags_str = ', '.join(set(common_tags_list)) if common_tags_list else '无'
                                        st.success(f"🎯 **记忆激活**: 发现与历史项目 `{', '.join(project_names)}` 存在关联（共同标签: {common_tags_str}），已进行横向对比。")
                                    else:
                                        st.info("🌱 **新物种收录**: 知识库中暂无同类，已作为种子存入大脑。")
                                    
                                    st.divider()
                                    
                                    # 显示清理后的 Markdown 内容（正文）
                                    cleaned_content = clean_markdown_for_display(result['markdown_content'])
                                    st.markdown(cleaned_content)
                                    
                                    word_buffer = result.get('word_buffer')
                                    if word_buffer:
                                        word_buffer.seek(0)
                                        st.download_button(
                                            label="📥 下载此报告",
                                            data=word_buffer.read(),
                                            file_name=f"{result['project_name']}_{datetime.now().strftime('%Y%m%d')}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key=f"download_{result['file_id']}"
                                        )
                        
                        # 保存结果到 session state
                        if "processed_results" not in st.session_state:
                            st.session_state.processed_results = []
                        st.session_state.processed_results.append(result)
                
                # 移动到下一个任务
                st.session_state.current_task_index += 1
                st.rerun()
            else:
                # 所有任务完成
                st.success(f"✅ 成功分析 {len(st.session_state.processed_results)}/{len(task_queue)} 个文件！")
                
                # 清理状态
                st.session_state.is_analyzing = False
                if "row_containers" in st.session_state:
                    del st.session_state.row_containers
                if "current_task_index" in st.session_state:
                    del st.session_state.current_task_index
                if "task_queue" in st.session_state:
                    del st.session_state.task_queue
                
                st.rerun()

    # 从 session state 渲染已处理的结果（刷新后仍然显示）
    if st.session_state.get("processed_results"):
        st.divider()
        st.subheader("📊 分析结果列表")
        
        selected_results = []
        for result in st.session_state["processed_results"]:
            is_selected = render_result_row(result, len(selected_results))
            if is_selected:
                selected_results.append(result)
        
        # 批量下载按钮
        st.divider()
        if selected_results:
            if st.button("📥 下载选中项目（ZIP 压缩包）", type="primary"):
                zip_buffer = io.BytesIO()
                
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for result in selected_results:
                        word_buffer = result.get('word_buffer')
                        if word_buffer:
                            word_buffer.seek(0)
                            file_name = f"{result['project_name']}_{result['file_id']}.docx"
                            zip_file.writestr(file_name, word_buffer.read())
                
                zip_buffer.seek(0)
                
                st.download_button(
                    label=f"⬇️ 点击下载 ZIP 文件（{len(selected_results)} 个报告）",
                    data=zip_buffer.read(),
                    file_name=f"batch_reports_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                    mime="application/zip",
                    key="download_zip"
                )
        
        # 清除所有结果按钮
        if st.button("🗑️ 清除所有结果"):
            st.session_state["processed_results"] = []
            st.session_state["task_queue"] = []
            st.session_state["processing_status"] = {}
            st.rerun()

elif page == "📂 全量知识库":
    # ========== 全量知识库页面 ==========
    st.title("📂 全量项目知识库")
    st.caption("查看和管理所有已分析的项目数据")
    
    try:
        # 加载知识库数据
        df_kb = load_knowledge_base()
        
        if df_kb.empty:
            st.info("📭 知识库为空，暂无项目数据。请先在工作台分析项目，数据会自动积累到这里。")
        else:
            # 格式化日期列
            if 'timestamp' in df_kb.columns:
                df_kb['timestamp'] = df_kb['timestamp'].apply(format_timestamp_for_display)
            
            # 计算统计指标
            stats = calculate_kb_statistics(df_kb)
            
            # 显示统计指标
            st.divider()
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("📊 已收录项目", stats['total_projects'])
            
            with col2:
                st.metric("⭐ 平均评分", f"{stats['avg_score']:.1f}" if stats['avg_score'] > 0 else "N/A")
            
            with col3:
                st.metric("🔥 最热赛道", stats['top_industry'])
            
            with col4:
                if stats['industry_count']:
                    top_count = list(stats['industry_count'].values())[0]
                    st.metric("🏆 最高赛道项目数", top_count)
                else:
                    st.metric("🏆 最高赛道项目数", 0)
            
            st.divider()
            
            # 显示全量数据表格
            st.subheader("📋 项目数据表")
            st.caption("💡 提示: 表格支持排序、搜索和筛选功能")
            
            # 选择要显示的列
            display_columns = ['timestamp', 'project_name', 'industry', 'tags', 'stage', 'score', 'summary', 'risk_level']
            available_columns = [col for col in display_columns if col in df_kb.columns]
            
            if available_columns:
                # 使用 st.dataframe 展示全量数据，开启排序和搜索
                st.dataframe(
                    df_kb[available_columns],
                    use_container_width=True,
                    hide_index=True,
                    height=600  # 增加表格高度以便查看更多数据
                )
            else:
                st.dataframe(
                    df_kb,
                    use_container_width=True,
                    hide_index=True,
                    height=600
                )
            
            # 导出功能
            st.divider()
            st.subheader("📥 数据导出")
            col1, col2 = st.columns(2)
            
            with col1:
                # 导出 CSV
                csv = df_kb.to_csv(index=False, encoding='utf-8-sig')
                st.download_button(
                    label="📥 下载 CSV 文件",
                    data=csv,
                    file_name=f"project_database_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime="text/csv",
                    help="下载完整项目数据库（CSV 格式，支持 Excel 打开）"
                )
            
            with col2:
                # 导出 JSON
                json_str = df_kb.to_json(orient='records', force_ascii=False, indent=2)
                st.download_button(
                    label="📥 下载 JSON 文件",
                    data=json_str,
                    file_name=f"project_database_{datetime.now().strftime('%Y%m%d')}.json",
                    mime="application/json",
                    help="下载完整项目数据库（JSON 格式）"
                )
            
            # 赛道分布分析
            if stats['industry_count'] and len(stats['industry_count']) > 0:
                st.divider()
                st.subheader("📈 赛道分布")
                industry_df = pd.DataFrame([
                    {'赛道': industry, '项目数': count}
                    for industry, count in stats['industry_count'].items()
                ]).sort_values('项目数', ascending=False)
                
                st.dataframe(
                    industry_df,
                    use_container_width=True,
                    hide_index=True
                )
    
    except Exception as e:
        st.error(f"❌ 加载知识库数据失败: {str(e)}")
        st.info("💡 提示: 请检查 project_database.csv 文件是否存在且格式正确")

elif page == "📜 历史记录":
    # ========== 历史记录页面 ==========
    st.title("📜 历史记录")
    st.caption("查看所有已保存的分析历史记录")
    
    try:
        history_entries = load_history_entries()
        
        # 显示知识库统计
        kb_count = memory_manager.get_count()
        
        if not history_entries:
            st.info("📭 暂无历史记录。分析项目后，记录会自动保存到这里。")
        else:
            col1, col2 = st.columns(2)
            with col1:
                st.metric("📚 总记录数", len(history_entries))
            with col2:
                st.metric("🧠 知识库已收录项目", kb_count)
            st.divider()
            
            # 显示历史记录列表
            for idx, entry in enumerate(history_entries):
                datetime_str = entry.get("datetime", entry.get("timestamp", ""))
                mode_name = entry.get("mode", "未知模式")
                file_name = entry.get("file_name", "未知文件")
                
                # 格式化显示时间
                try:
                    if "T" in datetime_str:
                        dt = datetime.fromisoformat(datetime_str)
                        display_time = dt.strftime("%Y-%m-%d %H:%M")
                    else:
                        display_time = datetime_str
                except:
                    display_time = datetime_str
                
                with st.expander(f"📄 {display_time} | {mode_name} | {file_name}", expanded=False):
                    col1, col2 = st.columns([1, 1])
                    
                    with col1:
                        st.caption(f"📋 **分析模式**: {mode_name}")
                        st.caption(f"📁 **文件名**: {file_name}")
                        st.caption(f"🕒 **分析时间**: {display_time}")
                    
                    with col2:
                        # 下载按钮
                        content = entry.get("content", "")
                        if content:
                            # 生成 Word 文档（临时）
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            temp_filename = f"{timestamp}_{Path(file_name).stem}.docx"
                            temp_path = HISTORY_DIR / temp_filename
                            
                            if markdown_to_docx(content, str(temp_path)):
                                with open(temp_path, "rb") as f:
                                    st.download_button(
                                        label="📥 下载 Word 报告",
                                        data=f.read(),
                                        file_name=temp_filename,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key=f"hist_download_{idx}"
                                    )
                                
                                # 清理临时文件
                                try:
                                    temp_path.unlink()
                                except:
                                    pass
                    
                    st.divider()
                    
                    # 显示分析内容（清理 JSON 代码块）
                    if content:
                        cleaned_content = clean_markdown_for_display(content)
                        st.markdown(cleaned_content)
                
                if idx < len(history_entries) - 1:
                    st.divider()
    
    except Exception as e:
        st.error(f"❌ 加载历史记录失败: {str(e)}")
        st.info("💡 提示: 请检查 history_data 文件夹是否存在且包含有效的 JSON 文件")
