"""
GP_Partner_Pro - 高级投资辅助系统（主动学习版本）
支持多模式切换、OpenRouter API、历史记录和报告导出
集成主动学习 Agent：RSS 抓取、网络搜索、LLM 总结、知识库检索
"""

import streamlit as st
import os
import json
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional, Dict, List, Union
import re
import zipfile
import io
import hashlib
import time

# 第三方库导入（必需依赖）
try:
    from openai import OpenAI
    from dotenv import load_dotenv
    import PyPDF2
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import pandas as pd
except ImportError as e:
    st.error(f"❌ 缺少必要的依赖库: {e}")
    st.stop()

# 立即加载环境变量（在导入其他模块之前）
# 确保在 Windows 和其他平台上都能正确读取项目根目录的 .env 文件
load_dotenv(dotenv_path=Path(__file__).parent / '.env', override=False)

# 软依赖导入（可选，用于主动学习功能）
HAS_DUCKDUCKGO = False
HAS_FEEDPARSER = False
HAS_SUPABASE = False
try:
    from duckduckgo_search import DDGS
    HAS_DUCKDUCKGO = True
except ImportError:
    pass

try:
    import feedparser
    HAS_FEEDPARSER = True
except ImportError:
    pass

try:
    from db.supabase_db import get_supabase_client
    HAS_SUPABASE = True
except ImportError:
    pass

# 注：已放弃向量检索方案，改用基于 LLM 标签提取的轻量级记忆系统
# 注意：环境变量已在上面加载（在导入 dotenv 之后立即调用）

# 页面配置
st.set_page_config(
    page_title="GP Partner Pro",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 常量定义
PROMPTS_DIR = Path("prompts")
HISTORY_DIR = Path("history_data")
BASE_URL = "https://openrouter.ai/api/v1"
DEFAULT_MODEL = "google/gemini-flash-1.5"
KNOWLEDGE_BASE_FILE = Path("project_database.csv")
MEMORY_STORE_FILE = Path("memory_store.json")  # 基于标签的记忆存储文件
EVOLUTION_LOG_FILE = Path("evolution_log.md")  # AI 进化日志文件
CONFIG_FILE = Path("config.json")  # 配置持久化文件
KNOWLEDGE_BRAIN_FILE = Path("knowledge_brain.json")  # 主动学习知识库文件

# 确保必要的文件夹存在
HISTORY_DIR.mkdir(exist_ok=True)
PROMPTS_DIR.mkdir(exist_ok=True)


# ==================== 配置持久化系统 ====================

def load_config() -> Dict:
    """加载配置文件"""
    try:
        if CONFIG_FILE.exists():
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                config = json.load(f)
                return config
        else:
            # 返回默认配置
            return {
                "api_key": os.getenv("OPENROUTER_API_KEY", ""),
                "base_url": BASE_URL,
                "model": DEFAULT_MODEL,
                "http_proxy": "http://127.0.0.1:7890"
            }
    except Exception as e:
        print(f"⚠️ 加载配置文件失败: {e}")
        return {
            "api_key": os.getenv("OPENROUTER_API_KEY", ""),
            "base_url": BASE_URL,
            "model": DEFAULT_MODEL,
            "http_proxy": "http://127.0.0.1:7890"
        }


def save_config():
    """保存配置到文件"""
    try:
        config = {
            "api_key": st.session_state.get("api_key_input", ""),
            "base_url": st.session_state.get("base_url_input", BASE_URL),
            "model": st.session_state.get("model_input", DEFAULT_MODEL),
            "http_proxy": st.session_state.get("http_proxy_input", "http://127.0.0.1:7890")
        }
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"⚠️ 保存配置文件失败: {e}")


# ==================== 基于 LLM 标签提取的轻量级记忆系统 ====================

def extract_tags_from_text(text: str, client: OpenAI, model: str) -> List[str]:
    """使用 LLM 从文本中提取项目标签"""
    try:
        tag_extraction_prompt = f"""请从以下项目文本中提取3-5个核心标签（关键词），用于项目分类和匹配。

文本内容：
{text[:1000]}

要求：
1. 标签应该是项目的核心技术、行业、商业模式等关键特征
2. 返回格式：纯文本，用逗号分隔，不要使用任何标记符号
3. 例如：AI, 医疗, B2B, SaaS, 多模态

标签："""

        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "你是一个专业的项目标签提取助手。只需返回标签，用逗号分隔，不要添加任何解释。"},
                {"role": "user", "content": tag_extraction_prompt}
            ],
            temperature=0.3,
            max_tokens=100
        )
        
        tags_text = response.choices[0].message.content.strip()
        # 解析标签
        tags = [tag.strip() for tag in tags_text.split(',') if tag.strip()]
        return tags[:5]  # 最多返回5个标签
    except Exception as e:
        print(f"⚠️ 标签提取失败: {e}")
        return []


class MemoryManager:
    """基于 LLM 标签提取的轻量级记忆系统管理器（单例模式）"""
    _instance = None
    _initialized = False
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super(MemoryManager, cls).__new__(cls)
        return cls._instance
    
    def __init__(self):
        if self._initialized:
            return
        
        self.memory_store_path = MEMORY_STORE_FILE
        self.memories = []
        self.enabled = True
        
        # 加载记忆库
        self._load_memories()
        self._initialized = True
    
    def _load_memories(self):
        """从 JSON 文件加载记忆库"""
        try:
            if self.memory_store_path.exists():
                with open(self.memory_store_path, 'r', encoding='utf-8') as f:
                    self.memories = json.load(f)
            else:
                self.memories = []
        except Exception as e:
            print(f"⚠️ 加载记忆库失败: {e}")
            self.memories = []
    
    def _save_memories(self):
        """保存记忆库到 JSON 文件"""
        try:
            with open(self.memory_store_path, 'w', encoding='utf-8') as f:
                json.dump(self.memories, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            print(f"⚠️ 保存记忆库失败: {e}")
            return False
    
    def query_similar(self, text: str, client: OpenAI, model: str, top_k: int = 3) -> List[Dict]:
        """基于标签匹配查询相似的历史项目"""
        if not self.enabled or len(self.memories) == 0:
            return []
        
        try:
            # 提取当前项目的标签
            query_tags = extract_tags_from_text(text, client, model)
            if not query_tags:
                return []
            
            # 计算每个历史项目的标签匹配度
            scored_projects = []
            for memory in self.memories:
                memory_tags = memory.get('tags', [])
                if not memory_tags:
                    continue
                
                # 计算标签交集数量（简单的标签匹配）
                common_tags = set(query_tags) & set(memory_tags)
                match_score = len(common_tags) / max(len(query_tags), len(memory_tags)) if max(len(query_tags), len(memory_tags)) > 0 else 0
                
                if match_score > 0:  # 至少有一个共同标签
                    scored_projects.append({
                        'name': memory.get('id', '未知项目'),
                        'score': memory.get('score', 'N/A'),
                        'summary': memory.get('summary', ''),
                        'match_score': match_score,
                        'common_tags': list(common_tags)
                    })
            
            # 按匹配度排序，返回前 top_k 个
            scored_projects.sort(key=lambda x: x['match_score'], reverse=True)
            return scored_projects[:top_k]
        except Exception as e:
            print(f"⚠️ 查询相似项目失败: {e}")
            return []
    
    def add_memory(self, name: str, summary: str, full_text: str, score, tags: List[str] = None, meta: Dict = None):
        """添加项目到记忆库"""
        if not self.enabled:
            return False
        
        try:
            # 如果未提供标签，使用空列表（标签会在外部通过 LLM 提取）
            if tags is None:
                tags = []
            
            memory_entry = {
                "id": name,
                "summary": summary,
                "score": score,
                "tags": tags,
                "timestamp": datetime.now().isoformat()
            }
            
            if meta:
                memory_entry.update(meta)
            
            # 检查是否已存在同名项目，如果存在则更新，否则添加
            existing_index = None
            for i, mem in enumerate(self.memories):
                if mem.get('id') == name:
                    existing_index = i
                    break
            
            if existing_index is not None:
                self.memories[existing_index] = memory_entry
            else:
                self.memories.append(memory_entry)
            
            # 保存到文件
            return self._save_memories()
        except Exception as e:
            print(f"⚠️ 添加记忆失败: {e}")
            return False
    
    def get_count(self) -> int:
        """获取知识库中的项目数量"""
        return len(self.memories) if self.enabled else 0


# 全局 MemoryManager 实例
memory_manager = MemoryManager()


# ==================== 主动学习 Agent - KnowledgeManager ====================

class KnowledgeManager:
    """主动学习知识管理器：RSS 抓取、搜索、LLM 总结、知识库检索"""
    
    def __init__(self):
        self.knowledge_file = KNOWLEDGE_BRAIN_FILE
        self.knowledge_base = []  # 存储学习到的知识条目
        self._load_knowledge()
    
    def _load_knowledge(self):
        """加载知识库"""
        try:
            if self.knowledge_file.exists():
                with open(self.knowledge_file, 'r', encoding='utf-8') as f:
                    self.knowledge_base = json.load(f)
            else:
                self.knowledge_base = []
        except Exception as e:
            print(f"⚠️ 加载知识库失败: {e}")
            self.knowledge_base = []
    
    def _save_knowledge(self):
        """保存知识库"""
        try:
            with open(self.knowledge_file, 'w', encoding='utf-8') as f:
                json.dump(self.knowledge_base, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            print(f"⚠️ 保存知识库失败: {e}")
            return False
    
    def fetch_rss(self, rss_url: str, max_items: int = 10) -> List[Dict]:
        """从 RSS 源抓取内容"""
        if not HAS_FEEDPARSER:
            return []
        
        try:
            feed = feedparser.parse(rss_url)
            items = []
            for entry in feed.entries[:max_items]:
                items.append({
                    'title': entry.get('title', ''),
                    'link': entry.get('link', ''),
                    'summary': entry.get('summary', ''),
                    'published': entry.get('published', ''),
                    'source': rss_url
                })
            return items
        except Exception as e:
            print(f"⚠️ RSS 抓取失败 ({rss_url}): {e}")
            return []
    
    def optimize_query_with_llm(self, query: str, client: OpenAI, model: str) -> str:
        """面向机构信源生成英文搜索词"""
        try:
            # 获取当前年份（用于默认）
            from datetime import datetime
            current_year = datetime.now().year
            next_year = current_year + 1
            
            prompt = f"""Translate this topic into a SINGLE, CONCISE English search query optimized for institutional sources (a16z, YC, TechCrunch, FT, Bloomberg).

User topic: {query}

CRITICAL RULES:
1. 3-6 key terms only.
2. MUST include year (e.g., 2025, 2026). If user didn't specify, use {current_year} or {next_year}.
3. Use ONE investment research keyword: "thesis" / "outlook" / "funding" / "investment trends" / "market landscape"
4. Focus on core entity + year + investment keyword (e.g., "Y Combinator AI investment trends 2025")
5. Output ONLY the query string, no explanations.

Examples:
- "Y Combinator AI 2025" → "Y Combinator AI investment trends 2025"
- "A16Z funding" → "a16z funding outlook {current_year}"
- "AI market" → "AI investment trends {current_year}"

Optimized search query:"""
            
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a search query optimizer for institutional investment research. Output ONLY the optimized English search query (3-6 key terms with year and investment keyword), no explanations."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=60
            )
            
            optimized = response.choices[0].message.content.strip()
            # 移除可能的引号
            optimized = optimized.strip('"').strip("'")
            
            # 验证：如果查询太短或没有年份，尝试手动添加
            words = optimized.split()
            has_year = any(word.isdigit() and len(word) == 4 for word in words)
            if not has_year and len(words) < 6:
                # 添加默认年份
                optimized = f"{optimized} {current_year}"
                words = optimized.split()
            
            # 额外验证：如果优化后的查询仍然太长（超过6个单词），手动截取前6个
            if len(words) > 6:
                optimized = " ".join(words[:6])
                print(f"⚠️ 查询过长，已截取前6个词: '{optimized}'")
            
            return optimized if optimized else query
        except Exception as e:
            print(f"⚠️ 查询优化失败: {e}")
            return query  # 失败时返回原查询
    
    def search_web(self, query: str, max_results: int = 5, proxy: str = None, client: OpenAI = None, model: str = None, region: str = "us-en") -> List[Dict]:
        """
        机构级搜索：优先新闻源，反向过滤中文内容农场
        """
        if not HAS_DUCKDUCKGO:
            return []
        
        # 关键修复：确保使用传入的 query（应该是英文），不再重新翻译
        optimized_query = query.strip() if query else ""
        
        # 如果提供了 LLM 客户端和模型，且查询可能是中文，才进行优化
        if client and model:
            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in optimized_query) if optimized_query else False
            if has_chinese or len(optimized_query) <= 2:
                print(f"⚠️ [WARNING] 检测到中文查询或异常查询: '{optimized_query}'，尝试二次翻译...")
                optimized_query = self.optimize_query_with_llm(query, client, model)
                print(f"🔍 查询优化: '{query}' -> '{optimized_query}'")
        
        # 验证查询有效性
        if not optimized_query or len(optimized_query.strip()) < 2:
            print(f"❌ [ERROR] 查询无效: '{optimized_query}'，无法搜索")
            return []
        
        # (1) 统一代理格式
        proxy_str = None
        if proxy and proxy.strip():
            proxy_str = proxy.strip()
            if not proxy_str.startswith("http"):
                proxy_str = f"http://{proxy_str}"
        
        # (2) 构造"反向站点过滤" query
        chinese_sites = [
            "-site:zhihu.com",
            "-site:baidu.com",
            "-site:zhidao.baidu.com",
            "-site:wenku.baidu.com",
            "-site:csdn.net",
            "-site:juejin.cn",
            "-site:weibo.com",
            "-site:bilibili.com"
        ]
        filtered_query = f"{optimized_query} {' '.join(chinese_sites)}"
        
        original_query_for_log = optimized_query
        print(f"🔎 [Search] 原始查询: '{original_query_for_log}'")
        print(f"🔎 [Search] 最终查询（带反向过滤）: '{filtered_query[:100]}...' (长度: {len(filtered_query)})")
        print(f"🔎 [Search] 配置: region={region}, max_results={max_results}, proxy={proxy_str}")
        
        def _normalize_result(r: Dict, query_used: str, source_type: str) -> Dict:
            """字段归一化：兼容不同版本的 duckduckgo_search"""
            return {
                'title': r.get('title', r.get('headline', '')),
                'url': r.get('href', r.get('url', r.get('link', ''))),
                'snippet': r.get('body', r.get('snippet', r.get('summary', ''))),
                'source': r.get('source', '') if source_type == 'news' else '',
                'date': r.get('date', '') if source_type == 'news' else '',
                'query': query_used,
                'original_query': query,
                'source_type': source_type
            }
        
        def _ddg_news(q: str) -> List[Dict]:
            """内部执行函数：新闻搜索"""
            try:
                with DDGS(proxy=proxy_str, timeout=30) as ddgs:
                    raw_results = list(ddgs.news(
                        keywords=q,
                        region=region,
                        safesearch='off',
                        max_results=max_results
                    ))
                    normalized = [_normalize_result(r, q, 'news') for r in raw_results]
                    print(f"📰 [News] 命中 {len(normalized)} 条新闻结果")
                    return normalized
            except Exception as inner_e:
                print(f"⚠️ [News] 新闻搜索失败: {inner_e}")
                return []
        
        def _ddg_text(q: str) -> List[Dict]:
            """内部执行函数：文本搜索"""
            try:
                with DDGS(proxy=proxy_str, timeout=30) as ddgs:
                    raw_results = list(ddgs.text(
                        keywords=q,
                        region=region,
                        safesearch='off',
                        max_results=max_results
                    ))
                    normalized = [_normalize_result(r, q, 'text') for r in raw_results]
                    print(f"📄 [Text] 命中 {len(normalized)} 条文本结果")
                    return normalized
            except Exception as inner_e:
                print(f"⚠️ [Text] 文本搜索失败: {inner_e}")
                return []

        # (3) 搜索策略：优先 news，空结果再 text
        results = []
        news_count = 0
        text_count = 0
        used_fallback = False
        
        # 第一轮：使用带反向过滤的完整查询
        results = _ddg_news(filtered_query)
        news_count = len(results)
        
        if not results:
            print("⚠️ [Strategy] 新闻搜索无结果，切换到文本搜索...")
            results = _ddg_text(filtered_query)
            text_count = len(results)
        
        # 降级重试机制：如果仍为空，简化查询后再次尝试
        if not results and len(optimized_query.split()) > 3:
            used_fallback = True
            simplified_query_base = " ".join(optimized_query.split()[:3])
            simplified_query = f"{simplified_query_base} {' '.join(chinese_sites)}"
            print(f"⚠️ [Fallback] 触发降级查询: '{simplified_query_base}' (带反向过滤)")
            
            results = _ddg_news(simplified_query)
            if results:
                news_count = len(results)
            else:
                results = _ddg_text(simplified_query)
                if results:
                    text_count = len(results)
        
        # 最终日志汇总
        print(f"✅ [Summary] 搜索完成: News={news_count}, Text={text_count}, 总计={len(results)}, 降级={used_fallback}")

        return results
    
    def summarize_with_llm(self, content: str, client: OpenAI, model: str, topic: str = "") -> Optional[str]:
        """使用 LLM 总结内容（角色扮演：高级投资分析师）"""
        try:
            # 新的角色扮演 System Prompt
            system_prompt = """You are an elite **Investment Analyst** at a top-tier AI Incubator. You report directly to the **Senior Investment Director**.
Your goal is to scan external information and extract high-value **Investment Signals**, **Market Alpha**, and **Strategic Threats**.

**Evaluation Criteria (The "VC Filter"):**
1. **Signal vs. Noise:** Ignore generic news, marketing fluff, or consumer-level tutorials. Focus on *industry shifts, funding dynamics, emerging tech stacks, and competitor moves*.
2. **Relevance:** If the content offers NO value to an investor (e.g., a basic tutorial on "how to use ChatGPT" when asking about "Agent Trends", or games like '风灵月影', unrelated entertainment, car models when asking about YC/Y Combinator), output exactly: `Irrelevant`.

**Tone:**
Professional, sharp, critical, and forward-looking. No filler words."""

            user_prompt = f"""**Task:**
Analyze the provided search result text related to the user's topic: '{topic}'.

**Content to Analyze:**
{content[:3000]}

**Output Format (If relevant):**
Write a concise **Investment Memo** (in Chinese) containing:
- **💡 核心情报 (Core Intelligence):** The key fact/news.
- **📉 赛道影响 (Market Impact):** How this affects the AI ecosystem/startups.
- **⚔️ 机会与风险 (Opportunities & Risks):** What should the Director pay attention to?
- **📝 一句话总结 (Key Takeaway):** 20 words max, punchy style.

**Important:** 
- If the content is NOT relevant to investment analysis (e.g., games, tutorials, unrelated content), output ONLY the string: `Irrelevant`
- If relevant, provide the Investment Memo in the format above."""

            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.5,  # 稍微提高温度，让输出更有创造性
                max_tokens=500  # 增加 token 限制，因为 Investment Memo 格式更详细
            )
            
            summary = response.choices[0].message.content.strip()
            
            # 检查是否被标记为不相关（无投资价值）
            if summary.lower() == 'irrelevant' or len(summary) < 10:
                return None  # 返回 None 表示内容不相关/无投资价值
            
            return summary
        except Exception as e:
            print(f"⚠️ LLM 总结失败: {e}")
            return None
    
    def add_knowledge(self, title: str, content: str, source: str, summary: str = "", tags: List[str] = None):
        """添加知识条目"""
        try:
            entry = {
                'id': f"{title}_{datetime.now().isoformat()}",
                'title': title,
                'content': content,
                'summary': summary,
                'source': source,
                'tags': tags or [],
                'timestamp': datetime.now().isoformat()
            }
            self.knowledge_base.append(entry)
            self._save_knowledge()
            return True
        except Exception as e:
            print(f"⚠️ 添加知识失败: {e}")
            return False
    
    def query_knowledge(self, query: str, top_k: int = 5) -> List[Dict]:
        """查询知识库（基于关键词匹配）"""
        if not self.knowledge_base:
            return []
        
        try:
            query_lower = query.lower()
            scored = []
            for entry in self.knowledge_base:
                score = 0
                # 标题匹配
                if query_lower in entry.get('title', '').lower():
                    score += 3
                # 内容匹配
                if query_lower in entry.get('content', '').lower():
                    score += 2
                # 标签匹配
                for tag in entry.get('tags', []):
                    if query_lower in tag.lower():
                        score += 1
                
                if score > 0:
                    scored.append({
                        'entry': entry,
                        'score': score
                    })
            
            # 按分数排序
            scored.sort(key=lambda x: x['score'], reverse=True)
            return [item['entry'] for item in scored[:top_k]]
        except Exception as e:
            print(f"⚠️ 查询知识库失败: {e}")
            return []
    
    def get_count(self) -> int:
        """获取知识库条目数量"""
        return len(self.knowledge_base)
    
    def clear_all(self) -> bool:
        """清空所有知识库条目"""
        try:
            self.knowledge_base = []
            return self._save_knowledge()
        except Exception as e:
            print(f"⚠️ 清空知识库失败: {e}")
            return False


# 全局 KnowledgeManager 实例
knowledge_manager = KnowledgeManager()


def load_prompt_files() -> Dict[str, str]:
    """扫描 prompts 文件夹，加载所有 .txt 文件"""
    prompt_files = {}
    if not PROMPTS_DIR.exists():
        return prompt_files
    
    txt_files = list(PROMPTS_DIR.glob("*.txt"))
    for file_path in sorted(txt_files):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read().strip()
                # 使用文件名（不含扩展名）作为键
                name = file_path.stem
                prompt_files[name] = content
        except Exception as e:
            st.warning(f"⚠️ 无法读取文件 {file_path.name}: {e}")
    
    return prompt_files


def extract_text_from_pdf(file) -> str:
    """从 PDF 文件中提取文本"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"❌ PDF 解析错误: {e}")
        return ""


def extract_text_from_docx(file) -> str:
    """从 DOCX 文件中提取文本"""
    try:
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"❌ DOCX 解析错误: {e}")
        return ""


def extract_text_from_uploaded_file(uploaded_file) -> Optional[str]:
    """根据文件类型提取文本"""
    file_extension = Path(uploaded_file.name).suffix.lower()
    
    if file_extension == ".pdf":
        return extract_text_from_pdf(uploaded_file)
    elif file_extension in [".docx", ".doc"]:
        return extract_text_from_docx(uploaded_file)
    else:
        st.error(f"❌ 不支持的文件格式: {file_extension}")
        return None


def save_history_entry(mode_name: str, file_name: str, analysis_content: str):
    """保存历史记录到 JSON 文件"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    entry = {
        "mode": mode_name,
        "file_name": file_name,
        "timestamp": timestamp,
        "datetime": datetime.now().isoformat(),
        "content": analysis_content
    }
    
    filename = f"{timestamp}_{mode_name.replace(' ', '_')}.json"
    filepath = HISTORY_DIR / filename
    
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(entry, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        st.error(f"❌ 保存历史记录失败: {e}")
        return False


def load_history_entries() -> List[Dict]:
    """加载所有历史记录"""
    entries = []
    if not HISTORY_DIR.exists():
        return entries
    
    json_files = sorted(HISTORY_DIR.glob("*.json"), reverse=True)
    for filepath in json_files:
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                entry = json.load(f)
                entries.append(entry)
        except Exception as e:
            st.warning(f"⚠️ 无法读取历史记录 {filepath.name}: {e}")
    
    return entries


def clean_markdown_for_display(markdown_text: str) -> str:
    """清理 Markdown 文本，移除 Tags 头部、JSON 代码块和进化建议（用于显示和导出）"""
    cleaned = markdown_text
    # 移除 Tags 头部
    cleaned = re.sub(r'---TAGS:\s*\[[^\]]+\]---\s*\n?', '', cleaned)
    # 移除 ```json ... ``` 代码块
    pattern = r'```json\s*[\s\S]*?```'
    cleaned = re.sub(pattern, '', cleaned, flags=re.DOTALL)
    # 移除进化建议章节
    cleaned = re.sub(r'##\s*🧬\s*进化建议\s*\n.*?(?=\n```json|\n```|$)', '', cleaned, flags=re.DOTALL)
    # 清理多余的空行
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()


def markdown_to_docx(markdown_text: str, output_path: str):
    """将 Markdown 文本转换为 Word 文档"""
    # 清理 JSON 代码块
    markdown_text = clean_markdown_for_display(markdown_text)
    
    doc = Document()
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # 空行
            doc.add_paragraph()
            continue
        
        # 检查标题层级
        if line.startswith('# '):
            # H1 标题
            heading = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            # H2 标题
            heading = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            # H3 标题
            heading = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            # H4 标题
            heading = doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            para = doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.startswith(('  - ', '  * ', '    - ', '    * ')):
            # 嵌套列表项
            indent_level = len(line) - len(line.lstrip())
            para = doc.add_paragraph(line.lstrip()[2:].strip(), style='List Bullet')
            para.paragraph_format.left_indent = Pt(indent_level * 12)
        elif re.match(r'^\d+\.\s+', line):
            # 有序列表
            para = doc.add_paragraph(line, style='List Number')
        elif line.startswith('**') and line.endswith('**'):
            # 粗体文本
            para = doc.add_paragraph()
            para.add_run(line[2:-2]).bold = True
        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            # 斜体文本
            para = doc.add_paragraph()
            para.add_run(line[1:-1]).italic = True
        else:
            # 普通段落 - 处理行内格式（粗体、斜体等）
            para = doc.add_paragraph()
            # 处理粗体 **text**
            if '**' in line:
                parts = re.split(r'(\*\*[^\*]+?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        para.add_run(part[2:-2]).bold = True
                    elif part.strip():
                        # 处理斜体 *text*
                        if '*' in part and not part.startswith('*'):
                            italic_parts = re.split(r'(\*[^\*]+\*)', part)
                            for italic_part in italic_parts:
                                if italic_part.startswith('*') and italic_part.endswith('*') and len(italic_part) > 2:
                                    para.add_run(italic_part[1:-1]).italic = True
                                elif italic_part.strip():
                                    para.add_run(italic_part)
                        else:
                            para.add_run(part)
            else:
                # 没有粗体标记，直接添加文本
                para.add_run(line)
    
    try:
        doc.save(output_path)
        return True
    except Exception as e:
        st.error(f"❌ 保存 Word 文档失败: {e}")
        return False


def enhance_system_prompt(base_prompt: str, similar_projects: List[Dict] = None) -> str:
    """增强系统提示词，要求 LLM 严格按照特定格式输出"""
    
    # 输出格式要求
    format_instruction = """

---
【严格输出格式要求】

你的输出必须严格按照以下三部分结构：

**第一部分 - 标签头部（必须）**：
在内容最开始，输出一行：
---TAGS: ["标签1", "标签2", "标签3", "标签4", "标签5"]---

**第二部分 - 分析正文**：
正常的 Markdown 格式分析报告，包含所有分析内容。

**第三部分 - 进化建议（必须）**：
在报告最后，必须包含一个章节：
## 🧬 进化建议
[针对本次分析，反思当前 System Prompt 的不足，并给出具体的优化指令建议。建议应该具体、可操作，例如："建议在提示词中增加对XX赛道的特殊关注"或"建议明确要求分析XX维度的风险"等]

**第四部分 - JSON 数据块（必须）**：
在进化建议之后，附带一个 JSON 数据块：

```json
{
  "project_name": "项目名称",
  "industry": "所属赛道/行业（如：AI Agent, 具身智能）",
  "tags": ["技术标签1", "技术标签2", "技术标签3"],
  "stage": "融资阶段（如：Angel, Pre-A, A轮, B轮等）",
  "score": 8,
  "summary": "一句话核心评价（50字以内）",
  "risk_level": "High/Medium/Low"
}
```

注意：
- project_name: 从商业计划书中提取的项目名称（字符串）
- industry: 所属赛道/行业（字符串，如：AI Agent, 具身智能, SaaS, 区块链等）
- tags: 技术标签列表（数组，如：["RAG", "LLM", "SaaS", "多模态"]），必须与第一部分的 TAGS 保持一致
- stage: 融资阶段（字符串，如：Angel, Pre-A, A轮, B轮等，如果未提及则填写 "未披露"）
- score: 投资推荐评分（整数，1-10分，10分为最高），必须在正文中明确显示
- summary: 一句话核心评价（字符串，50字以内，简洁概括项目价值和风险）
- risk_level: 风险等级（字符串，必须是 "High", "Medium", "Low" 之一）

请严格按照以上格式输出，不要遗漏任何部分。
"""
    
    # 如果有相似项目，添加历史项目参考信息
    if similar_projects and len(similar_projects) > 0:
        similar_info = "\n\n---\n【历史项目参考】知识库中发现了与本项目相似的过往项目（基于标签匹配）：\n"
        for proj in similar_projects:
            common_tags_str = ', '.join(proj.get('common_tags', [])) if proj.get('common_tags') else '无'
            match_score = proj.get('match_score', 0)
            similar_info += f"- {proj['name']} (评分: {proj['score']}, 核心评价: {proj['summary']}, 共同标签: {common_tags_str}, 匹配度: {match_score:.2%})\n"
        similar_info += "\n请在分析时横向对比，指出本项目的差异化优势或重复造轮子的风险。参考这些历史案例，但不要被其局限，重点分析当前项目的独特价值。\n"
        return base_prompt + similar_info + format_instruction
    
    return base_prompt + format_instruction


def extract_tags_from_response(response_text: str) -> List[str]:
    """从 LLM 响应中提取 Tags（从头部 ---TAGS: ... --- 格式）"""
    try:
        # 匹配 ---TAGS: ["标签1", "标签2"]--- 格式
        pattern = r'---TAGS:\s*(\[[^\]]+\])---'
        match = re.search(pattern, response_text)
        
        if match:
            tags_str = match.group(1)
            # 解析 JSON 数组
            tags = json.loads(tags_str)
            if isinstance(tags, list):
                return [str(tag).strip() for tag in tags if tag]
        
        # 如果没有找到，尝试从 JSON 数据块中提取
        json_data = extract_json_from_response(response_text)
        if json_data and 'tags' in json_data:
            tags = json_data.get('tags', [])
            if isinstance(tags, list):
                return [str(tag).strip() for tag in tags if tag]
        
        return []
    except Exception as e:
        print(f"⚠️ 提取 Tags 失败: {e}")
        return []


def extract_evolution_suggestion(response_text: str) -> Optional[str]:
    """从 LLM 响应中提取进化建议（## 🧬 进化建议 后的内容）"""
    try:
        # 匹配 ## 🧬 进化建议 后的内容（直到 JSON 代码块或文件末尾）
        pattern = r'##\s*🧬\s*进化建议\s*\n(.*?)(?=\n```json|\n```|$)'
        match = re.search(pattern, response_text, re.DOTALL)
        
        if match:
            suggestion = match.group(1).strip()
            # 清理多余的空行
            suggestion = re.sub(r'\n{3,}', '\n\n', suggestion)
            return suggestion if suggestion else None
        
        return None
    except Exception as e:
        print(f"⚠️ 提取进化建议失败: {e}")
        return None


def extract_score_enhanced(response_text: str, json_data: Optional[Dict] = None) -> Optional[Union[int, float, str]]:
    """增强的分数提取函数，支持多种格式"""
    # 首先尝试从 JSON 数据中提取
    if json_data and 'score' in json_data:
        score = json_data.get('score')
        if score and score != 'N/A':
            try:
                if isinstance(score, (int, float)):
                    return score
                score_str = str(score).strip()
                # 尝试转换为数字
                if score_str.replace('.', '').isdigit():
                    return float(score_str) if '.' in score_str else int(score_str)
            except:
                pass
    
    # 在前 1000 字符中搜索分数
    search_text = response_text[:1000]
    
    # 模式1: Score: 8 或 评分：8.5
    patterns = [
        r'(?:Score|评分|分数|投资评分)[:：]\s*(\d+(?:\.\d+)?)',
        r'\[(\d+(?:\.\d+)?)分\]',
        r'评分[：:]\s*(\d+(?:\.\d+)?)',
        r'(\d+(?:\.\d+)?)\s*分',
        r'(\d+(?:\.\d+)?)\s*/\s*10',  # x/10 格式
    ]
    
    for pattern in patterns:
        match = re.search(pattern, search_text, re.IGNORECASE)
        if match:
            try:
                score_val = float(match.group(1))
                if 1 <= score_val <= 10:
                    return int(score_val) if score_val.is_integer() else score_val
            except:
                continue
    
    # 在整个文本中搜索 x/10 格式（兜底策略）
    x_10_pattern = r'(\d+(?:\.\d+)?)\s*/\s*10'
    matches = re.findall(x_10_pattern, response_text)
    if matches:
        for match in matches:
            try:
                score_val = float(match)
                if 1 <= score_val <= 10:
                    return int(score_val) if score_val.is_integer() else score_val
            except:
                continue
    
    return None


def parse_llm_response(response_text: str) -> Dict:
    """统一解析 LLM 响应，提取所有结构化信息"""
    result = {
        'tags': [],
        'json_data': None,
        'score': None,
        'evolution_suggestion': None,
        'body_content': response_text  # 原始内容
    }
    
    # 提取 Tags
    result['tags'] = extract_tags_from_response(response_text)
    
    # 提取 JSON 数据
    result['json_data'] = extract_json_from_response(response_text)
    
    # 提取分数（增强版）
    result['score'] = extract_score_enhanced(response_text, result['json_data'])
    
    # 提取进化建议
    result['evolution_suggestion'] = extract_evolution_suggestion(response_text)
    
    # 提取正文内容（移除 Tags 头部、进化建议和 JSON）
    body = response_text
    # 移除 Tags 头部
    body = re.sub(r'---TAGS:\s*\[[^\]]+\]---\s*\n?', '', body)
    # 移除进化建议章节
    body = re.sub(r'##\s*🧬\s*进化建议\s*\n.*?(?=\n```json|\n```|$)', '', body, flags=re.DOTALL)
    # 移除 JSON 代码块
    body = re.sub(r'```json\s*[\s\S]*?```', '', body)
    result['body_content'] = body.strip()
    
    return result


def save_evolution_suggestion(project_name: str, suggestion: str):
    """保存进化建议到日志文件"""
    try:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_entry = f"\n\n---\n## {timestamp} - {project_name}\n\n{suggestion}\n"
        
        with open(EVOLUTION_LOG_FILE, 'a', encoding='utf-8') as f:
            f.write(log_entry)
        return True
    except Exception as e:
        print(f"⚠️ 保存进化建议失败: {e}")
        return False


def extract_json_from_response(response_text: str) -> Optional[Dict]:
    """从 AI 响应中提取 JSON 数据块"""
    try:
        # 首先尝试匹配 ```json ... ``` 代码块（支持多行）
        pattern = r'```json\s*([\s\S]*?)\s*```'
        matches = re.findall(pattern, response_text)
        
        if matches:
            # 取最后一个匹配（通常 JSON 在末尾）
            json_str = matches[-1].strip()
            try:
                # 解析 JSON
                json_data = json.loads(json_str)
                # 验证是否包含必需字段
                required_fields = ['project_name', 'industry', 'tags', 'stage', 'score', 'summary', 'risk_level']
                if all(field in json_data for field in required_fields):
                    return json_data
            except json.JSONDecodeError:
                # JSON 解析失败，尝试清理后再解析
                # 移除可能的注释或额外字符
                json_str = re.sub(r'//.*', '', json_str)  # 移除单行注释
                json_str = re.sub(r'/\*.*?\*/', '', json_str, flags=re.DOTALL)  # 移除多行注释
                try:
                    json_data = json.loads(json_str)
                    required_fields = ['project_name', 'industry', 'tags', 'stage', 'score', 'summary', 'risk_level']
                    if all(field in json_data for field in required_fields):
                        return json_data
                except:
                    pass
        
        # 如果没有找到代码块，尝试直接查找 JSON 对象（兜底方案）
        # 使用更精确的正则表达式匹配完整的 JSON 对象
        json_pattern = r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}'
        json_matches = re.findall(json_pattern, response_text, re.DOTALL)
        
        if json_matches:
            # 从后往前尝试解析（通常 JSON 在末尾）
            for json_str in reversed(json_matches):
                try:
                    json_data = json.loads(json_str)
                    # 验证是否包含必需字段
                    required_fields = ['project_name', 'industry', 'tags', 'stage', 'score', 'summary', 'risk_level']
                    if all(field in json_data for field in required_fields):
                        return json_data
                except:
                    continue
        
        return None
    except Exception as e:
        # 静默失败，不中断主流程
        print(f"⚠️ 提取 JSON 失败: {e}")
        return None


def save_to_knowledge_base(json_data: Dict):
    """将提取的 JSON 数据保存到项目知识库 CSV 文件"""
    try:
        # 添加时间戳
        json_data['timestamp'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 将 tags 列表转换为字符串（CSV 格式，用分号分隔）
        if isinstance(json_data.get('tags'), list):
            json_data['tags'] = ';'.join(json_data['tags'])
        
        # 定义 CSV 列顺序
        columns = ['timestamp', 'project_name', 'industry', 'tags', 'stage', 'score', 'summary', 'risk_level']
        
        # 确保所有字段都存在
        for col in columns:
            if col not in json_data:
                json_data[col] = ''
        
        # 创建 DataFrame
        df_new = pd.DataFrame([json_data])
        
        # 检查文件是否存在
        if KNOWLEDGE_BASE_FILE.exists():
            # 读取现有数据
            df_existing = pd.read_csv(KNOWLEDGE_BASE_FILE, encoding='utf-8-sig')
            # 合并数据
            df_combined = pd.concat([df_existing, df_new], ignore_index=True)
        else:
            # 创建新文件
            df_combined = df_new
        
        # 按照列顺序重新排列
        df_combined = df_combined[columns]
        
        # 保存到 CSV（使用 utf-8-sig 编码以支持中文，避免 Windows Excel 乱码）
        df_combined.to_csv(KNOWLEDGE_BASE_FILE, index=False, encoding='utf-8-sig')
        return True
    except Exception as e:
        # 静默失败，不中断主流程
        print(f"⚠️ 保存知识库失败: {e}")
        return False


def load_knowledge_base() -> pd.DataFrame:
    """加载项目知识库"""
    try:
        if KNOWLEDGE_BASE_FILE.exists():
            df = pd.read_csv(KNOWLEDGE_BASE_FILE, encoding='utf-8-sig')
            # 确保 timestamp 列为字符串类型，便于格式化
            if 'timestamp' in df.columns:
                df['timestamp'] = df['timestamp'].astype(str)
            return df
        else:
            return pd.DataFrame()
    except Exception as e:
        print(f"⚠️ 加载知识库失败: {e}")
        return pd.DataFrame()


def format_timestamp_for_display(timestamp_str: str) -> str:
    """将时间戳格式化为 YYYY-MM-DD 格式"""
    try:
        # 尝试解析多种时间格式
        if isinstance(timestamp_str, str):
            # 处理 "YYYY-MM-DD HH:MM:SS" 格式
            if ' ' in timestamp_str:
                date_part = timestamp_str.split(' ')[0]
                return date_part
            # 处理 ISO 格式
            elif 'T' in timestamp_str:
                date_part = timestamp_str.split('T')[0]
                return date_part
            # 如果已经是 YYYY-MM-DD 格式
            elif len(timestamp_str) >= 10:
                return timestamp_str[:10]
        return timestamp_str
    except:
        return timestamp_str


def get_recent_projects(limit: int = 5) -> List[Dict]:
    """获取最近的项目记录"""
    try:
        df = load_knowledge_base()
        if df.empty:
            return []
        
        # 按时间倒序排列，取最新 N 条
        if 'timestamp' in df.columns and 'project_name' in df.columns:
            df_sorted = df.sort_values('timestamp', ascending=False)
            recent = df_sorted.head(limit)
            
            result = []
            for _, row in recent.iterrows():
                date_str = format_timestamp_for_display(str(row.get('timestamp', '')))
                project_name = str(row.get('project_name', '未知项目'))
                result.append({
                    'date': date_str,
                    'name': project_name
                })
            return result
        return []
    except Exception as e:
        print(f"⚠️ 获取最近项目失败: {e}")
        return []


def calculate_kb_statistics(df: pd.DataFrame) -> Dict:
    """计算知识库统计指标"""
    stats = {
        'total_projects': 0,
        'avg_score': 0.0,
        'top_industry': 'N/A',
        'industry_count': {}
    }
    
    try:
        if df.empty:
            return stats
        
        stats['total_projects'] = len(df)
        
        # 计算平均评分
        if 'score' in df.columns:
            try:
                scores = pd.to_numeric(df['score'], errors='coerce').dropna()
                if len(scores) > 0:
                    stats['avg_score'] = round(scores.mean(), 1)
            except:
                pass
        
        # 计算最热赛道
        if 'industry' in df.columns:
            industry_counts = df['industry'].value_counts()
            if len(industry_counts) > 0:
                stats['top_industry'] = industry_counts.index[0]
                stats['industry_count'] = industry_counts.to_dict()
        
        return stats
    except Exception as e:
        print(f"⚠️ 计算统计指标失败: {e}")
        return stats


def generate_file_id(file_name: str) -> str:
    """为文件生成唯一 ID"""
    # 使用文件名和当前时间戳生成唯一 ID
    content = f"{file_name}_{time.time()}"
    return hashlib.md5(content.encode()).hexdigest()[:12]


def process_single_file(
    uploaded_file,
    file_id: str,
    system_prompt: str,
    api_key: str,
    model: str,
    selected_mode: str,
    row_container,
    similar_projects: List[Dict] = None
) -> Optional[Dict]:
    """处理单个文件，返回结果字典"""
    try:
        # 提取文件文本
        file_text = extract_text_from_uploaded_file(uploaded_file)
        if not file_text:
            with row_container.container():
                st.error(f"❌ 无法从文件 {uploaded_file.name} 中提取文本")
            return None
        
        # 初始化 OpenAI 客户端（使用侧边栏配置的值）
        # 注意：base_url 和 api_key 应该从侧边栏的输入框获取
        # 这里需要从调用函数的参数中获取，或从 session_state 读取
        client = OpenAI(
            base_url=BASE_URL,  # Base URL 是固定的，使用常量
            api_key=api_key
        )
        
        # system_prompt 已经在外部增强（包含 RAG 信息），直接使用
        enhanced_prompt = system_prompt
        
        # 调用 API（流式）
        stream = call_openrouter_api(client, enhanced_prompt, file_text, model)
        if not stream:
            with row_container.container():
                st.error(f"❌ API 调用失败: {uploaded_file.name}")
            return None
        
        # 收集完整响应
        full_response = ""
        
        for chunk in stream:
            if chunk.choices[0].delta.content is not None:
                full_response += chunk.choices[0].delta.content
        
        # 解析 LLM 响应（提取 Tags、JSON、分数、进化建议等）
        parsed = parse_llm_response(full_response)
        json_data = parsed.get('json_data')
        extracted_tags = parsed.get('tags', [])
        extracted_score = parsed.get('score')
        evolution_suggestion = parsed.get('evolution_suggestion')
        body_content = parsed.get('body_content', full_response)
        
        # 如果 JSON 中有 tags，优先使用 JSON 中的（更准确）
        if json_data and json_data.get('tags'):
            extracted_tags = json_data.get('tags', [])
        
        # 如果 JSON 中有 score，优先使用 JSON 中的
        if json_data and json_data.get('score') and json_data.get('score') != 'N/A':
            extracted_score = json_data.get('score')
        
        # 保存到 CSV 知识库
        if json_data:
            save_to_knowledge_base(json_data)
        
        # 保存历史记录（保存完整内容）
        save_history_entry(selected_mode, uploaded_file.name, full_response)
        
        # 保存进化建议
        if evolution_suggestion:
            project_name = json_data.get('project_name', uploaded_file.name) if json_data else uploaded_file.name
            save_evolution_suggestion(project_name, evolution_suggestion)
        
        # 生成 Word 文档（内存中，使用清理后的正文内容）
        word_buffer = io.BytesIO()
        temp_path = HISTORY_DIR / f"temp_{file_id}.docx"
        # Word 导出时使用清理后的正文（不包含 Tags 头部和进化建议）
        markdown_to_docx(body_content, str(temp_path))
        
        with open(temp_path, "rb") as f:
            word_buffer.write(f.read())
        word_buffer.seek(0)
        
        # 清理临时文件
        try:
            temp_path.unlink()
        except:
            pass
        
        # 返回结果字典
        project_name = json_data.get('project_name', uploaded_file.name) if json_data else uploaded_file.name
        final_score = extracted_score if extracted_score is not None else (json_data.get('score', 'N/A') if json_data else 'N/A')
        
        return {
            'file_id': file_id,
            'file_name': uploaded_file.name,
            'markdown_content': body_content,  # 使用清理后的正文
            'full_response': full_response,  # 保留完整响应用于展示
            'word_buffer': word_buffer,
            'json_data': json_data,
            'project_name': project_name,
            'score': final_score,
            'tags': extracted_tags,
            'evolution_suggestion': evolution_suggestion,
            'similar_projects': similar_projects if similar_projects else [],
            'status': 'completed',
            'timestamp': datetime.now().isoformat()
        }
    
    except Exception as e:
        row_container.error(f"❌ 处理文件 {uploaded_file.name} 时出错: {str(e)}")
        return None


def render_result_row(result: Dict, index: int) -> bool:
    """渲染单个结果行，返回是否被选中"""
    col1, col2 = st.columns([0.05, 0.95])
    
    with col1:
        checkbox_key = f"select_{result['file_id']}"
        is_selected = st.checkbox(
            "",
            value=True,
            key=checkbox_key,
            label_visibility="collapsed"
        )
    
    with col2:
        # 构建 expander 标题
        score_display = f"{result['score']}分" if isinstance(result['score'], (int, float)) else str(result['score'])
        expander_title = f"✅ [{score_display}] {result['project_name']} (点击展开详情)"
        
        with st.expander(expander_title, expanded=False):
            # ========== 大脑思考路径展示 ==========
            st.info("🧠 **大脑思考路径**")
            
            # 显示提取到的标签
            tags = result.get('tags', [])
            if tags:
                tags_str = ', '.join([f"`{tag}`" for tag in tags])
                st.markdown(f"**核心关键词**: {tags_str}")
            else:
                st.markdown("**核心关键词**: 未提取到标签")
            
            # 显示记忆激活状态
            similar_projects = result.get('similar_projects', [])
            if similar_projects and len(similar_projects) > 0:
                project_names = [proj.get('name', '未知项目') for proj in similar_projects]
                common_tags_list = []
                for proj in similar_projects:
                    common_tags = proj.get('common_tags', [])
                    if common_tags:
                        common_tags_list.extend(common_tags)
                
                common_tags_str = ', '.join(set(common_tags_list)) if common_tags_list else '无'
                st.success(f"🎯 **记忆激活**: 发现与历史项目 `{', '.join(project_names)}` 存在关联（共同标签: {common_tags_str}），已进行横向对比。")
            else:
                st.info("🌱 **新物种收录**: 知识库中暂无同类，已作为种子存入大脑。")
            
            st.divider()
            
            # 显示清理后的 Markdown 内容（正文）
            cleaned_content = clean_markdown_for_display(result['markdown_content'])
            st.markdown(cleaned_content)
            
            # 单独下载按钮
            word_buffer = result.get('word_buffer')
            if word_buffer:
                word_buffer.seek(0)
                st.download_button(
                    label="📥 下载此报告",
                    data=word_buffer.read(),
                    file_name=f"{result['project_name']}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_row_{result['file_id']}"
                )
    
    return is_selected


def call_openrouter_api(client: OpenAI, system_prompt: str, user_content: str, model: str):
    """调用 OpenRouter API 进行流式响应"""
    try:
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content}
        ]
        
        stream = client.chat.completions.create(
            model=model,
            messages=messages,
            stream=True
        )
        
        return stream
    except Exception as e:
        st.error(f"❌ API 调用错误: {e}")
        return None


# ==================== 侧边栏 ====================
with st.sidebar:
    st.header("⚙️ 控制台")
    
    # 初始化 is_analyzing 状态
    if "is_analyzing" not in st.session_state:
        st.session_state.is_analyzing = False
    
    # 页面导航（如果正在分析则禁用）
    page = st.sidebar.radio(
        "页面导航",
        options=["🚀 AI 分析工作台", "🧠 知识大脑", "📂 全量知识库", "📜 历史记录", "📡 雷达候选池", "📄 周报"],
        index=0,
        label_visibility="visible",
        disabled=st.session_state.is_analyzing
    )
    
    # 如果正在分析，显示警告
    if st.session_state.is_analyzing:
        st.warning("⚠️ 分析任务运行中，导航已锁定...")
    
    st.divider()
    
    # API 配置（持久化）
    st.subheader("API 配置")
    
    # 加载配置（仅在第一次时）
    if "config_loaded" not in st.session_state:
        config = load_config()
        st.session_state["api_key_input"] = config.get("api_key", os.getenv("OPENROUTER_API_KEY", ""))
        st.session_state["base_url_input"] = config.get("base_url", BASE_URL)
        st.session_state["model_input"] = config.get("model", DEFAULT_MODEL)
        st.session_state["http_proxy_input"] = config.get("http_proxy", "http://127.0.0.1:7890")
        st.session_state["config_loaded"] = True
        
        # 如果配置文件不存在，首次加载时保存默认配置
        if not CONFIG_FILE.exists():
            save_config()
    
    # Base URL（禁用，固定值，但保存在配置中）
    base_url = st.text_input(
        "Base URL",
        value=st.session_state.get("base_url_input", BASE_URL),
        disabled=True,
        key="base_url_input",
        help="OpenRouter API 基础 URL（固定）"
    )
    
    # API Key（支持自动保存）
    api_key = st.text_input(
        "API Key",
        value=st.session_state.get("api_key_input", ""),
        type="password",
        key="api_key_input",
        on_change=save_config,
        help="从配置文件加载或手动输入，修改后自动保存"
    )
    
    # Model（支持自动保存）
    model = st.text_input(
        "Model",
        value=st.session_state.get("model_input", DEFAULT_MODEL),
        key="model_input",
        on_change=save_config,
        help="OpenRouter 模型名称，修改后自动保存"
    )
    
    # HTTP Proxy（支持自动保存，用于网络搜索）
    http_proxy = st.text_input(
        "HTTP Proxy",
        value=st.session_state.get("http_proxy_input", "http://127.0.0.1:7890"),
        key="http_proxy_input",
        on_change=save_config,
        help="HTTP 代理地址（用于网络搜索，留空则不使用代理）。默认: http://127.0.0.1:7890",
        placeholder="http://127.0.0.1:7890"
    )
    
    # 显示配置状态提示
    if CONFIG_FILE.exists():
        st.caption("💾 配置已自动保存")
    
    st.divider()
    
    # 分析模式选择（仅在工作台页面需要，且在未分析时显示）
    if page == "🚀 AI 分析工作台" and not st.session_state.is_analyzing:
        st.subheader("📋 分析模式")
        prompt_files = load_prompt_files()
        
        if not prompt_files:
            st.warning("⚠️ 请在 prompts 文件夹放入 .txt 提示词文件")
            selected_mode = None
            system_prompt = None
        else:
            mode_names = list(prompt_files.keys())
            selected_mode = st.selectbox(
                "选择分析模式",
                options=mode_names,
                index=0
            )
            base_prompt = prompt_files.get(selected_mode, "")
            # 存储基础提示词，稍后在调用 API 时增强
            system_prompt = base_prompt
    else:
        # 知识库页面不需要这些变量，但需要初始化以避免错误
        # 如果正在分析，从 session_state 恢复
        if st.session_state.is_analyzing and "selected_mode" in st.session_state:
            selected_mode = st.session_state.selected_mode
            system_prompt = st.session_state.get("system_prompt", "")
        else:
            selected_mode = None
            system_prompt = None
    
    st.divider()
    
    # 调试模式开关
    debug_mode = st.checkbox("🔧 调试模式", value=st.session_state.get("debug_mode", False), key="debug_mode_checkbox")
    st.session_state["debug_mode"] = debug_mode
    
    # 调试面板（仅在调试模式开启时显示）
    if st.session_state.get("debug_mode", False):
        st.divider()
        st.subheader("🔧 调试面板")
        
        # Supabase 连接状态
        st.markdown("**Supabase 连接状态**")
        supabase_url = os.getenv("SUPABASE_URL")
        supabase_anon_key = os.getenv("SUPABASE_ANON_KEY")
        
        if not supabase_url:
            st.error("❌ 未设置 SUPABASE_URL")
        else:
            st.success(f"✅ SUPABASE_URL: {supabase_url[:30]}..." if len(supabase_url) > 30 else f"✅ SUPABASE_URL: {supabase_url}")
        
        if not supabase_anon_key:
            st.error("❌ 未设置 SUPABASE_ANON_KEY")
        else:
            masked_key = supabase_anon_key[:6] + "***" if len(supabase_anon_key) > 6 else "***"
            st.success(f"✅ SUPABASE_ANON_KEY: {masked_key}")
        
        # 检查 Supabase 客户端
        if HAS_SUPABASE:
            try:
                test_client = get_supabase_client(use_service_role=False)
                if test_client:
                    st.success("✅ Supabase 客户端已创建")
                else:
                    st.error("❌ Supabase 客户端创建失败")
            except Exception as e:
                st.error(f"❌ Supabase 客户端错误: {e}")
        else:
            st.error("❌ Supabase 库未安装")
        
        st.divider()
        
        # Deals 读取状态（仅在"雷达候选池"页面显示）
        if page == "📡 雷达候选池":
            st.markdown("**Deals 读取状态**")
            deals_count = st.session_state.get("debug_deals_count")
            filtered_count = st.session_state.get("debug_filtered_deals_count")
            
            if deals_count is not None:
                st.info(f"📊 读取条数: {deals_count}")
                if filtered_count is not None:
                    st.info(f"🔍 过滤后条数: {filtered_count}")
            else:
                st.caption("等待页面加载...")
        
        # Weekly Reports 读取状态（仅在"周报"页面显示）
        if page == "📄 周报":
            st.markdown("**Weekly Reports 读取状态**")
            weekly_report_debug = st.session_state.get("debug_weekly_report")
            
            if weekly_report_debug == "no reports":
                st.warning("⚠️ no reports")
            elif isinstance(weekly_report_debug, dict):
                st.success(f"✅ week_start: {weekly_report_debug.get('week_start', 'N/A')}")
                st.info(f"📄 markdown 长度: {weekly_report_debug.get('markdown_length', 0)} 字符")
            else:
                st.caption("等待页面加载...")
        
        # Deal Actions 最近一次写入
        if page == "📡 雷达候选池":
            st.divider()
            st.markdown("**Deal Actions 最近一次写入**")
            last_payload = st.session_state.get("last_action_payload")
            
            if last_payload:
                st.json(last_payload)
            else:
                st.caption("暂无写入记录")
    
    st.divider()
    
    # AI 进化日志
    with st.expander("🧬 查看 AI 进化日志", expanded=False):
        try:
            if EVOLUTION_LOG_FILE.exists():
                with open(EVOLUTION_LOG_FILE, 'r', encoding='utf-8') as f:
                    log_content = f.read()
                if log_content.strip():
                    st.markdown(log_content)
                else:
                    st.info("暂无进化建议，分析项目后会自动生成。")
            else:
                st.info("暂无进化建议，分析项目后会自动生成。")
        except Exception as e:
            st.warning(f"⚠️ 读取进化日志失败: {e}")


# ==================== 主界面 ====================
# 根据页面导航显示不同内容
if page == "🚀 AI 分析工作台":
    # ========== AI 分析工作台页面 ==========
    st.title("🚀 AI 分析工作台")
    st.caption("高级投资辅助系统 - 支持批量分析与报告导出")

    # 显示当前模式
    if selected_mode:
        st.info(f"📋 **当前分析模式**: {selected_mode}")
    else:
        st.warning("⚠️ 请先在侧边栏选择分析模式")

    st.divider()

    # 初始化 session state
    if "processed_results" not in st.session_state:
        st.session_state["processed_results"] = []
    if "task_queue" not in st.session_state:
        st.session_state["task_queue"] = []
    if "processing_status" not in st.session_state:
        st.session_state["processing_status"] = {}

    # 初始化 MemoryManager（轻量级，无需额外依赖）
    # memory_manager 已在模块加载时自动初始化

    # 批量文件上传（分析中时禁用）
    uploaded_files = st.file_uploader(
        "📄 上传分析文件（支持批量）",
        type=["pdf", "docx", "doc"],
        help="支持 PDF 和 DOCX 格式，可同时上传多个文件进行批量分析",
        accept_multiple_files=True,
        disabled=st.session_state.is_analyzing
    )

    # 开始批量分析按钮（仅在未分析时显示）
    if uploaded_files and len(uploaded_files) > 0 and not st.session_state.is_analyzing:
        if st.button("🚀 开始批量分析", type="primary", disabled=not (selected_mode and api_key)):
            if not api_key:
                st.error("❌ 请输入 API Key")
                st.stop()
            
            if not selected_mode:
                st.error("❌ 请选择分析模式")
                st.stop()
            
            # 设置分析状态并保存配置到 session_state
            st.session_state.is_analyzing = True
            st.session_state.selected_mode = selected_mode
            st.session_state.system_prompt = system_prompt
            
            # 初始化任务队列
            task_queue = []
            for i, uploaded_file in enumerate(uploaded_files):
                file_id = generate_file_id(f"{uploaded_file.name}_{i}_{time.time()}")
                task_queue.append({
                    'file_id': file_id,
                    'file': uploaded_file,
                    'index': i
                })
            
            st.session_state.task_queue = task_queue
            st.rerun()
    
    # 如果正在分析，执行批量处理循环
    if st.session_state.is_analyzing and "task_queue" in st.session_state:
        task_queue = st.session_state.task_queue
        if len(task_queue) > 0:
            st.subheader("📋 任务队列")
            
            # 初始化所有行的占位符（仅在第一次渲染时）
            if "row_containers" not in st.session_state:
                st.session_state.row_containers = []
                for task in task_queue:
                    row_container = st.empty()
                    with row_container.container():
                        st.markdown(f"⏳ **等待处理**: {task['file'].name}")
                    st.session_state.row_containers.append({
                        'file_id': task['file_id'],
                        'container': row_container,
                        'status': 'waiting'
                    })
            
            # 获取当前处理索引
            if "current_task_index" not in st.session_state:
                st.session_state.current_task_index = 0
            
            current_idx = st.session_state.current_task_index
            
            if current_idx < len(task_queue):
                task = task_queue[current_idx]
                row_info = st.session_state.row_containers[current_idx]
                
                # 更新当前行的状态
                with row_info['container'].container():
                    st.markdown(f"🔄 **正在深入分析 (关联知识库中...)**: {task['file'].name}...")
                
                # 提取文件文本
                file_text = extract_text_from_uploaded_file(task['file'])
                
                if file_text:
                    # 初始化 OpenAI 客户端（用于标签提取和相似项目查询）
                    client = OpenAI(
                        base_url=base_url,  # 使用侧边栏配置的值
                        api_key=api_key
                    )
                    
                    # 基于标签匹配检索相似项目
                    similar_projects = []
                    if memory_manager.enabled and memory_manager.get_count() > 0:
                        similar_projects = memory_manager.query_similar(file_text, client, model, top_k=3)
                    
                    # 检索主动学习知识库（如果知识库不为空）
                    knowledge_context = ""
                    if knowledge_manager.get_count() > 0:
                        try:
                            # 从文件文本中提取关键词用于检索
                            # 简单提取：取前 200 字符作为查询
                            query_text = file_text[:200]
                            knowledge_entries = knowledge_manager.query_knowledge(query_text, top_k=3)
                            
                            if knowledge_entries:
                                knowledge_context = "\n\n---\n【主动学习知识库参考】以下是从最新行业动态中学习到的相关信息：\n"
                                for entry in knowledge_entries:
                                    knowledge_context += f"- {entry.get('title', '无标题')}：{entry.get('summary', '无总结')}\n"
                                knowledge_context += "\n请结合这些行业动态，分析当前项目的市场时机和竞争环境。\n"
                        except Exception as e:
                            print(f"⚠️ 知识库检索失败: {e}")
                            # 静默失败，不影响主流程
                    
                    # 增强 Prompt（包含相似项目信息和知识库上下文）
                    enhanced_prompt = enhance_system_prompt(
                        st.session_state.system_prompt,
                        similar_projects if similar_projects else None
                    )
                    
                    # 如果有知识库上下文，添加到 prompt
                    if knowledge_context:
                        enhanced_prompt = enhanced_prompt + knowledge_context
                    
                    # 处理文件
                    result = process_single_file(
                        uploaded_file=task['file'],
                        file_id=task['file_id'],
                        system_prompt=enhanced_prompt,
                        api_key=api_key,
                        model=model,
                        selected_mode=st.session_state.selected_mode,
                        row_container=row_info['container'],
                        similar_projects=similar_projects
                    )
                    
                    if result:
                        # 提取 JSON 数据
                        json_data = result.get('json_data')
                        
                        # 存储到轻量级记忆库
                        if json_data and memory_manager.enabled:
                            project_name = json_data.get('project_name', task['file'].name)
                            summary = json_data.get('summary', '')
                            score = result.get('score', json_data.get('score', 'N/A'))
                            # 使用从响应中提取的 tags（已通过 parse_llm_response 提取）
                            tags = result.get('tags', [])
                            
                            # 如果仍然没有标签，尝试从 JSON 中获取
                            if not tags and json_data.get('tags'):
                                tags = json_data.get('tags', [])
                            
                            memory_manager.add_memory(
                                name=project_name,
                                summary=summary,
                                full_text=result.get('full_response', result['markdown_content']),  # 使用完整响应
                                score=score,
                                tags=tags,
                                meta={
                                    'industry': json_data.get('industry', ''),
                                    'stage': json_data.get('stage', ''),
                                    'risk_level': json_data.get('risk_level', '')
                                }
                            )
                        
                        # 更新行容器为最终结果
                        row_info['container'].empty()
                        with row_info['container'].container():
                            col1, col2 = st.columns([0.05, 0.95])
                            
                            with col1:
                                checkbox_key = f"select_{result['file_id']}"
                                st.checkbox(
                                    "",
                                    value=True,
                                    key=checkbox_key,
                                    label_visibility="collapsed"
                                )
                            
                            with col2:
                                score_display = f"{result['score']}分" if isinstance(result['score'], (int, float)) else str(result['score'])
                                expander_title = f"✅ [{score_display}] {result['project_name']} (点击展开详情)"
                                
                                with st.expander(expander_title, expanded=False):
                                    # ========== 大脑思考路径展示 ==========
                                    st.info("🧠 **大脑思考路径**")
                                    
                                    # 显示提取到的标签
                                    tags = result.get('tags', [])
                                    if tags:
                                        tags_str = ', '.join([f"`{tag}`" for tag in tags])
                                        st.markdown(f"**核心关键词**: {tags_str}")
                                    else:
                                        st.markdown("**核心关键词**: 未提取到标签")
                                    
                                    # 显示记忆激活状态
                                    similar_projects = result.get('similar_projects', [])
                                    if similar_projects and len(similar_projects) > 0:
                                        project_names = [proj.get('name', '未知项目') for proj in similar_projects]
                                        common_tags_list = []
                                        for proj in similar_projects:
                                            common_tags = proj.get('common_tags', [])
                                            if common_tags:
                                                common_tags_list.extend(common_tags)
                                        
                                        common_tags_str = ', '.join(set(common_tags_list)) if common_tags_list else '无'
                                        st.success(f"🎯 **记忆激活**: 发现与历史项目 `{', '.join(project_names)}` 存在关联（共同标签: {common_tags_str}），已进行横向对比。")
                                    else:
                                        st.info("🌱 **新物种收录**: 知识库中暂无同类，已作为种子存入大脑。")
                                    
                                    st.divider()
                                    
                                    # 显示清理后的 Markdown 内容（正文）
                                    cleaned_content = clean_markdown_for_display(result['markdown_content'])
                                    st.markdown(cleaned_content)
                                    
                                    word_buffer = result.get('word_buffer')
                                    if word_buffer:
                                        word_buffer.seek(0)
                                        st.download_button(
                                            label="📥 下载此报告",
                                            data=word_buffer.read(),
                                            file_name=f"{result['project_name']}_{datetime.now().strftime('%Y%m%d')}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key=f"download_{result['file_id']}"
                                        )
                        
                        # 保存结果到 session state
                        if "processed_results" not in st.session_state:
                            st.session_state.processed_results = []
                        st.session_state.processed_results.append(result)
                
                # 移动到下一个任务
                st.session_state.current_task_index += 1
                st.rerun()
            else:
                # 所有任务完成
                st.success(f"✅ 成功分析 {len(st.session_state.processed_results)}/{len(task_queue)} 个文件！")
                
                # 清理状态
                st.session_state.is_analyzing = False
                if "row_containers" in st.session_state:
                    del st.session_state.row_containers
                if "current_task_index" in st.session_state:
                    del st.session_state.current_task_index
                if "task_queue" in st.session_state:
                    del st.session_state.task_queue
                
                st.rerun()

    # 从 session state 渲染已处理的结果（刷新后仍然显示）
    if st.session_state.get("processed_results"):
        st.divider()
        st.subheader("📊 分析结果列表")
        
        selected_results = []
        for result in st.session_state["processed_results"]:
            is_selected = render_result_row(result, len(selected_results))
            if is_selected:
                selected_results.append(result)
        
        # 批量下载按钮
        st.divider()
        if selected_results:
            if st.button("📥 下载选中项目（ZIP 压缩包）", type="primary"):
                zip_buffer = io.BytesIO()
                
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for result in selected_results:
                        word_buffer = result.get('word_buffer')
                        if word_buffer:
                            word_buffer.seek(0)
                            file_name = f"{result['project_name']}_{result['file_id']}.docx"
                            zip_file.writestr(file_name, word_buffer.read())
                
                zip_buffer.seek(0)
                
                st.download_button(
                    label=f"⬇️ 点击下载 ZIP 文件（{len(selected_results)} 个报告）",
                    data=zip_buffer.read(),
                    file_name=f"batch_reports_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                    mime="application/zip",
                    key="download_zip"
                )
        
        # 清除所有结果按钮
        if st.button("🗑️ 清除所有结果"):
            st.session_state["processed_results"] = []
            st.session_state["task_queue"] = []
            st.session_state["processing_status"] = {}
            st.rerun()

elif page == "🧠 知识大脑":
    # ========== 知识大脑页面 ==========
    st.title("🧠 知识大脑")
    st.caption("主动学习 Agent - RSS 抓取、网络搜索、LLM 总结、知识库检索")
    
    # 检查依赖状态
    if not HAS_DUCKDUCKGO or not HAS_FEEDPARSER:
        st.warning("⚠️ **软依赖未安装**：请安装依赖以启用联网功能")
        st.info("安装命令：`pip install duckduckgo-search feedparser`")
        st.caption("💡 基础分析功能仍然可用，但主动学习功能将被禁用")
        st.divider()
    
    # 知识库状态
    kb_count = knowledge_manager.get_count()
    st.metric("📚 知识库条目数", kb_count)
    
    st.divider()
    
    # 启动专项学习
    st.subheader("🚀 启动专项学习")
    
    col1, col2 = st.columns(2)
    
    with col1:
        learning_mode = st.radio(
            "学习模式",
            options=["RSS 源学习", "网络搜索学习"],
            disabled=not (HAS_DUCKDUCKGO and HAS_FEEDPARSER)
        )
    
    with col2:
        topic = st.text_input(
            "学习主题/关键词",
            placeholder="例如：AI Agent, 具身智能, SaaS...",
            disabled=not (HAS_DUCKDUCKGO and HAS_FEEDPARSER)
        )
    
    # RSS 源配置
    if learning_mode == "RSS 源学习":
        st.subheader("📡 RSS 源配置")
        rss_sources = st.text_area(
            "RSS 源列表（每行一个 URL）",
            placeholder="https://example.com/rss1\nhttps://example.com/rss2",
            disabled=not HAS_FEEDPARSER,
            height=100
        )
        
        max_items = st.number_input("每个源最多抓取条目数", min_value=1, max_value=50, value=10, disabled=not HAS_FEEDPARSER)
        
        if st.button("📥 开始 RSS 学习", type="primary", disabled=not (HAS_FEEDPARSER and rss_sources and topic)):
            if not api_key:
                st.error("❌ 请输入 API Key")
            else:
                with st.spinner("正在抓取 RSS 并总结..."):
                    client = OpenAI(base_url=base_url, api_key=api_key)
                    rss_urls = [url.strip() for url in rss_sources.split('\n') if url.strip()]
                    
                    total_fetched = 0
                    total_summarized = 0
                    total_filtered = 0
                    
                    for rss_url in rss_urls:
                        try:
                            items = knowledge_manager.fetch_rss(rss_url, max_items)
                            total_fetched += len(items)
                            
                            for item in items:
                                content = f"{item.get('title', '')}\n{item.get('summary', '')}"
                                summary = knowledge_manager.summarize_with_llm(
                                    content, client, model, topic
                                )
                                
                                # 只有当 summary 不为 None（内容相关）时才保存
                                if summary:
                                    knowledge_manager.add_knowledge(
                                        title=item.get('title', '无标题'),
                                        content=content,
                                        source=rss_url,
                                        summary=summary,
                                        tags=[topic] if topic else []
                                    )
                                    total_summarized += 1
                                else:
                                    total_filtered += 1
                        except Exception as e:
                            st.warning(f"⚠️ 处理 RSS 源失败 ({rss_url}): {e}")
                    
                    if total_summarized > 0:
                        st.success(f"✅ 完成！抓取 {total_fetched} 条，通过验证并保存 {total_summarized} 条，过滤掉 {total_filtered} 条无关内容")
                    else:
                        st.warning(f"⚠️ 未找到有效信息！抓取 {total_fetched} 条，但所有结果都被过滤为不相关内容。请尝试更精确的主题关键词。")
                    st.rerun()
    
    # 网络搜索学习
    elif learning_mode == "网络搜索学习":
        st.subheader("🔍 网络搜索学习")
        
        search_query = st.text_input(
            "搜索关键词",
            value=topic if topic else "",
            placeholder="例如：AI Agent 投资趋势",
            disabled=not HAS_DUCKDUCKGO
        )
        
        max_results = st.number_input("最多搜索结果数", min_value=1, max_value=20, value=5, disabled=not HAS_DUCKDUCKGO)
        
        if st.button("🔍 开始搜索学习", type="primary", disabled=not (HAS_DUCKDUCKGO and search_query)):
            if not api_key:
                st.error("❌ 请输入 API Key")
            else:
                # 使用 st.status 显示详细进度
                with st.status("正在进行深度学习...", expanded=True) as status:
                    client = OpenAI(base_url=base_url, api_key=api_key)
                    
                    try:
                        # 步骤1: 强制翻译为英文搜索词（关键修复）
                        status.write("🔄 步骤1/4: 正在将关键词转化为英文...")
                        
                        # 强制调用 LLM 翻译函数，确保使用英文搜索
                        english_query = knowledge_manager.optimize_query_with_llm(
                            search_query, 
                            client, 
                            model
                        )
                        
                        # 关键修复：验证翻译结果，确保是有效的英文查询
                        if not english_query or not english_query.strip():
                            english_query = search_query  # 如果翻译失败，使用原查询
                            status.write(f"⚠️ 翻译失败，使用原始查询: {search_query}")
                        else:
                            english_query = english_query.strip()
                            
                            # 验证：如果翻译后的查询仍然是中文（或只有一个字符），强制重新翻译
                            # 检查是否包含中文字符
                            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in english_query)
                            if has_chinese or len(english_query) <= 2:
                                status.write(f"⚠️ 翻译结果异常（可能仍是中文），强制重新翻译...")
                                # 强制使用更直接的翻译 Prompt
                                try:
                                    translation_prompt = f"Translate this to a concise English search query. Input: '{search_query}'. Output: ONLY the English query (3-5 words max)."
                                    translation_response = client.chat.completions.create(
                                        model=model,
                                        messages=[
                                            {"role": "system", "content": "You are a search query translator. Output ONLY the English search query, no explanations."},
                                            {"role": "user", "content": translation_prompt}
                                        ],
                                        temperature=0.3,
                                        max_tokens=30
                                    )
                                    english_query = translation_response.choices[0].message.content.strip()
                                    english_query = english_query.strip('"').strip("'")
                                    # 如果还是有问题，至少确保不是单字符
                                    if len(english_query) <= 2:
                                        status.write(f"❌ 翻译严重失败，使用原查询作为备选")
                                        english_query = search_query
                                except Exception as e:
                                    status.write(f"⚠️ 强制翻译失败: {e}，使用原查询")
                                    english_query = search_query
                        
                        # UI 反馈：显示翻译结果
                        status.write(f"🎯 已锁定英文搜索词: **{english_query}**")
                        status.write(f"📝 原始查询: {search_query} → 英文查询: {english_query}")
                        
                        # 获取代理配置
                        proxy_config = st.session_state.get("http_proxy_input", "http://127.0.0.1:7890")
                        # 如果代理为空字符串，设为 None
                        proxy_config = proxy_config.strip() if proxy_config and proxy_config.strip() else None
                        
                        # 步骤2: 使用英文查询进行搜索（强制美国区，优先新闻源）
                        status.write("🔍 步骤2/4: 正在全球网络搜索 (Region: US-EN, 优先新闻源)...")
                        
                        # 添加调试日志，确认传入的参数
                        print(f"[DEBUG] 调用 search_web，参数 query={english_query}, max_results={max_results}")
                        status.write(f"🎯 搜索关键词: '{english_query}'")
                        
                        results = knowledge_manager.search_web(
                            english_query,  # 强制使用翻译后的英文查询
                            max_results, 
                            proxy=proxy_config,
                            client=None,  # 不再需要在这里翻译，因为已经翻译完成
                            model=None,
                            region="us-en"  # 强制美国区
                        )
                        
                        # 如果没有结果，可能是网络连接问题
                        if not results:
                            status.write("❌ 网络连接错误：未搜索到结果")
                            st.error("网络连接错误：未搜索到结果。请检查 1. 代理端口是否正确; 2. VPN是否开启。")
                        else:
                            # 统计新闻和文本结果数量
                            news_results = [r for r in results if r.get('source_type') == 'news']
                            text_results = [r for r in results if r.get('source_type') == 'text']
                            status.write(f"✅ 成功搜索到 {len(results)} 条结果 (新闻: {len(news_results)}, 文本: {len(text_results)})")
                            
                            # 显示结果来源域名（去重）
                            domains = set()
                            for r in results:
                                url = r.get('url', '')
                                if url:
                                    try:
                                        from urllib.parse import urlparse
                                        domain = urlparse(url).netloc
                                        if domain:
                                            domains.add(domain)
                                    except:
                                        pass
                            if domains:
                                status.write(f"📊 来源域名: {', '.join(list(domains)[:5])}{'...' if len(domains) > 5 else ''}")
                            
                            status.write("📖 步骤3/4: 正在分析并提取知识...")
                            
                            total_summarized = 0
                            total_filtered = 0
                            
                            # 逐条处理搜索结果
                            for i, result in enumerate(results):
                                title = result.get('title', '无标题')
                                url = result.get('url', result.get('link', ''))
                                
                                # 提取域名用于显示
                                domain = ""
                                if url:
                                    try:
                                        from urllib.parse import urlparse
                                        domain = urlparse(url).netloc
                                    except:
                                        pass
                                
                                domain_display = f" [{domain}]" if domain else ""
                                status.write(f"📄 正在阅读第 {i+1}/{len(results)} 篇: {title[:40]}{domain_display}...")
                                
                                content = f"{result.get('title', '')}\n{result.get('snippet', '')}"
                                summary = knowledge_manager.summarize_with_llm(
                                    content, client, model, search_query
                                )
                                
                                # 显式展示 LLM 的判断
                                if summary:
                                    # 成功总结
                                    print(f"[✅ 成功] 成功提取见解: {title} (来源: {domain})")
                                    status.write(f"✅ [{i+1}] 成功提取见解: {title[:40]}{domain_display}")
                                    
                                    # 保存到知识库（使用归一化后的 url 字段）
                                    success = knowledge_manager.add_knowledge(
                                        title=title,
                                        content=content,
                                        source=url,  # 使用归一化后的 url
                                        summary=summary,
                                        tags=[search_query] if search_query else []
                                    )
                                    
                                    if success:
                                        total_summarized += 1
                                        # 使用 toast 提示（如果 Streamlit 版本支持）
                                        try:
                                            st.toast(f"✅ 已保存: {title[:30]}...", icon="✅")
                                        except:
                                            pass  # 如果 toast 不可用，静默跳过
                                    else:
                                        status.write(f"⚠️ [{i+1}] 保存失败: {title[:50]}")
                                else:
                                    # 被过滤为不相关
                                    print(f"[过滤] 判定为无关: {title} (来源: {domain})")
                                    status.write(f"⏭️ [{i+1}] 跳过: {title[:40]} (原因: 无投资价值/不相关){domain_display}")
                                    total_filtered += 1
                                    # 使用 toast 提示
                                    try:
                                        st.toast(f"⏭️ 跳过: {title[:30]}...（无关内容）", icon="⏭️")
                                    except:
                                        pass  # 如果 toast 不可用，静默跳过
                            
                            # 步骤4: 完成
                            status.write("💾 步骤4/4: 保存知识库...")
                            
                            # 确保保存到文件（add_knowledge 内部已经保存，但再次确认）
                            knowledge_manager._save_knowledge()
                            
                            # 显示最终结果
                            if total_summarized > 0:
                                status.write(f"✅ 完成！共处理 {len(results)} 条，成功保存 {total_summarized} 条，过滤 {total_filtered} 条")
                                st.success(f"✅ 完成！共处理 {len(results)} 条，成功保存 {total_summarized} 条，过滤 {total_filtered} 条")
                            else:
                                status.write(f"⚠️ 未找到有效信息！所有结果都被过滤为不相关内容")
                                st.warning(f"⚠️ 未找到有效信息！搜索 {len(results)} 条，但所有结果都被过滤为不相关内容。请尝试更精确的搜索关键词。")
                        
                    except Exception as e:
                        error_msg = f"❌ 搜索学习失败: {str(e)}"
                        print(error_msg)
                        try:
                            status.write(error_msg)
                        except:
                            pass  # 如果 status 已失效，静默跳过
                        st.error(f"{error_msg}。请检查 1. 代理端口是否正确; 2. VPN是否开启; 3. API Key 是否有效。")
                    
                    # 必须刷新页面，让用户立刻看到新知识（无论成功还是失败）
                    st.rerun()
    
    st.divider()
    
    # 知识库浏览
    col_header1, col_header2 = st.columns([0.9, 0.1])
    with col_header1:
        st.subheader("📚 知识库浏览")
    with col_header2:
        # 清空知识库按钮
        if kb_count > 0:
            if st.button("🗑️ 清空", type="secondary", help="清空所有已学习知识（不可恢复）"):
                if knowledge_manager.clear_all():
                    st.success("✅ 知识库已清空")
                    st.rerun()
                else:
                    st.error("❌ 清空失败")
    
    if kb_count == 0:
        st.info("📭 知识库为空，请先启动专项学习")
    else:
        # 搜索框
        search_query_kb = st.text_input("🔍 搜索知识库", placeholder="输入关键词搜索...")
        
        if search_query_kb:
            results = knowledge_manager.query_knowledge(search_query_kb, top_k=10)
            if results:
                st.success(f"找到 {len(results)} 条相关条目")
                for idx, entry in enumerate(results):
                    with st.expander(f"📄 {entry.get('title', '无标题')}", expanded=False):
                        st.caption(f"来源: {entry.get('source', '未知')}")
                        st.caption(f"时间: {entry.get('timestamp', '')}")
                        if entry.get('tags'):
                            st.caption(f"标签: {', '.join(entry['tags'])}")
                        st.markdown(f"**总结**: {entry.get('summary', '无总结')}")
                        if entry.get('content'):
                            with st.expander("查看完整内容", expanded=False):
                                st.text(entry['content'])
            else:
                st.info("未找到相关条目")
        else:
            # 显示最近的知识条目
            st.caption("最近学习到的知识（最新 10 条）")
            recent_entries = list(reversed(knowledge_manager.knowledge_base[-10:]))
            for idx, entry in enumerate(recent_entries):
                with st.expander(f"📄 {entry.get('title', '无标题')}", expanded=False):
                    st.caption(f"来源: {entry.get('source', '未知')}")
                    st.caption(f"时间: {entry.get('timestamp', '')}")
                    if entry.get('tags'):
                        st.caption(f"标签: {', '.join(entry['tags'])}")
                    st.markdown(f"**总结**: {entry.get('summary', '无总结')}")
                    if entry.get('content'):
                        with st.expander("查看完整内容", expanded=False):
                            st.text(entry['content'])

elif page == "📂 全量知识库":
    # ========== 全量知识库页面 ==========
    st.title("📂 全量项目知识库")
    st.caption("查看和管理所有已分析的项目数据")
    
    try:
        # 加载知识库数据
        df_kb = load_knowledge_base()
        
        if df_kb.empty:
            st.info("📭 知识库为空，暂无项目数据。请先在工作台分析项目，数据会自动积累到这里。")
        else:
            # 格式化日期列
            if 'timestamp' in df_kb.columns:
                df_kb['timestamp'] = df_kb['timestamp'].apply(format_timestamp_for_display)
            
            # 计算统计指标
            stats = calculate_kb_statistics(df_kb)
            
            # 显示统计指标
            st.divider()
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("📊 已收录项目", stats['total_projects'])
            
            with col2:
                st.metric("⭐ 平均评分", f"{stats['avg_score']:.1f}" if stats['avg_score'] > 0 else "N/A")
            
            with col3:
                st.metric("🔥 最热赛道", stats['top_industry'])
            
            with col4:
                if stats['industry_count']:
                    top_count = list(stats['industry_count'].values())[0]
                    st.metric("🏆 最高赛道项目数", top_count)
                else:
                    st.metric("🏆 最高赛道项目数", 0)
            
            st.divider()
            
            # 显示全量数据表格
            st.subheader("📋 项目数据表")
            st.caption("💡 提示: 表格支持排序、搜索和筛选功能")
            
            # 选择要显示的列
            display_columns = ['timestamp', 'project_name', 'industry', 'tags', 'stage', 'score', 'summary', 'risk_level']
            available_columns = [col for col in display_columns if col in df_kb.columns]
            
            if available_columns:
                # 使用 st.dataframe 展示全量数据，开启排序和搜索
                st.dataframe(
                    df_kb[available_columns],
                    use_container_width=True,
                    hide_index=True,
                    height=600  # 增加表格高度以便查看更多数据
                )
            else:
                st.dataframe(
                    df_kb,
                    use_container_width=True,
                    hide_index=True,
                    height=600
                )
            
            # 导出功能
            st.divider()
            st.subheader("📥 数据导出")
            col1, col2 = st.columns(2)
            
            with col1:
                # 导出 CSV
                csv = df_kb.to_csv(index=False, encoding='utf-8-sig')
                st.download_button(
                    label="📥 下载 CSV 文件",
                    data=csv,
                    file_name=f"project_database_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime="text/csv",
                    help="下载完整项目数据库（CSV 格式，支持 Excel 打开）"
                )
            
            with col2:
                # 导出 JSON
                json_str = df_kb.to_json(orient='records', force_ascii=False, indent=2)
                st.download_button(
                    label="📥 下载 JSON 文件",
                    data=json_str,
                    file_name=f"project_database_{datetime.now().strftime('%Y%m%d')}.json",
                    mime="application/json",
                    help="下载完整项目数据库（JSON 格式）"
                )
            
            # 赛道分布分析
            if stats['industry_count'] and len(stats['industry_count']) > 0:
                st.divider()
                st.subheader("📈 赛道分布")
                industry_df = pd.DataFrame([
                    {'赛道': industry, '项目数': count}
                    for industry, count in stats['industry_count'].items()
                ]).sort_values('项目数', ascending=False)
                
                st.dataframe(
                    industry_df,
                    use_container_width=True,
                    hide_index=True
                )
    
    except Exception as e:
        st.error(f"❌ 加载知识库数据失败: {str(e)}")
        st.info("💡 提示: 请检查 project_database.csv 文件是否存在且格式正确")

elif page == "📜 历史记录":
    # ========== 历史记录页面 ==========
    st.title("📜 历史记录")
    st.caption("查看所有已保存的分析历史记录")
    
    try:
        history_entries = load_history_entries()
        
        # 显示知识库统计
        kb_count = memory_manager.get_count()
        
        if not history_entries:
            st.info("📭 暂无历史记录。分析项目后，记录会自动保存到这里。")
        else:
            col1, col2 = st.columns(2)
            with col1:
                st.metric("📚 总记录数", len(history_entries))
            with col2:
                st.metric("🧠 知识库已收录项目", kb_count)
            st.divider()
            
            # 显示历史记录列表
            for idx, entry in enumerate(history_entries):
                datetime_str = entry.get("datetime", entry.get("timestamp", ""))
                mode_name = entry.get("mode", "未知模式")
                file_name = entry.get("file_name", "未知文件")
                
                # 格式化显示时间
                try:
                    if "T" in datetime_str:
                        dt = datetime.fromisoformat(datetime_str)
                        display_time = dt.strftime("%Y-%m-%d %H:%M")
                    else:
                        display_time = datetime_str
                except:
                    display_time = datetime_str
                
                with st.expander(f"📄 {display_time} | {mode_name} | {file_name}", expanded=False):
                    col1, col2 = st.columns([1, 1])
                    
                    with col1:
                        st.caption(f"📋 **分析模式**: {mode_name}")
                        st.caption(f"📁 **文件名**: {file_name}")
                        st.caption(f"🕒 **分析时间**: {display_time}")
                    
                    with col2:
                        # 下载按钮
                        content = entry.get("content", "")
                        if content:
                            # 生成 Word 文档（临时）
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            temp_filename = f"{timestamp}_{Path(file_name).stem}.docx"
                            temp_path = HISTORY_DIR / temp_filename
                            
                            if markdown_to_docx(content, str(temp_path)):
                                with open(temp_path, "rb") as f:
                                    st.download_button(
                                        label="📥 下载 Word 报告",
                                        data=f.read(),
                                        file_name=temp_filename,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key=f"hist_download_{idx}"
                                    )
                                
                                # 清理临时文件
                                try:
                                    temp_path.unlink()
                                except:
                                    pass
                    
                    st.divider()
                    
                    # 显示分析内容（清理 JSON 代码块）
                    if content:
                        cleaned_content = clean_markdown_for_display(content)
                        st.markdown(cleaned_content)
                
                if idx < len(history_entries) - 1:
                    st.divider()
    
    except Exception as e:
        st.error(f"❌ 加载历史记录失败: {str(e)}")
        st.info("💡 提示: 请检查 history_data 文件夹是否存在且包含有效的 JSON 文件")

elif page == "📡 雷达候选池":
    # ========== 雷达候选池页面 ==========
    st.title("📡 雷达候选池")
    st.caption("查看和管理从雷达抓取中发现的潜在投资项目")
    
    if not HAS_SUPABASE:
        st.error("❌ Supabase 客户端未安装。请运行: `pip install supabase`")
        st.stop()
    
    # 获取 Supabase 客户端（前端只使用 ANON_KEY）
    supabase_client = get_supabase_client(use_service_role=False)
    
    if not supabase_client:
        st.error("❌ Supabase 连接失败。请检查环境变量：")
        st.code("""
SUPABASE_URL=your_supabase_url
SUPABASE_ANON_KEY=your_supabase_anon_key
        """)
        st.info("💡 请在 `.env` 文件中设置这些环境变量，或使用系统环境变量。")
        st.stop()
    
    # 初始化 session state
    if "deals_filter" not in st.session_state:
        st.session_state.deals_filter = "全部"
    if "deals_search" not in st.session_state:
        st.session_state.deals_search = ""
    
    # 过滤和搜索选项
    col_filter, col_search = st.columns([1, 2])
    with col_filter:
        filter_option = st.selectbox(
            "筛选条件",
            options=["全部", "未处理", "已标记"],
            index=0 if st.session_state.deals_filter == "全部" else (1 if st.session_state.deals_filter == "未处理" else 2),
            key="deals_filter_select"
        )
        st.session_state.deals_filter = filter_option
    
    with col_search:
        search_query = st.text_input(
            "🔍 关键字搜索",
            value=st.session_state.deals_search,
            placeholder="搜索项目名称或简介...",
            key="deals_search_input"
        )
        st.session_state.deals_search = search_query
    
    st.divider()
    
    try:
        # 先获取所有 deals（按 updated_at 倒序，限制 50 条）
        # 明确指定需要的字段（至少 id,title,canonical_name,one_liner,hostname,website,url,tags,score,created_at）
        response = supabase_client.table("deals")\
            .select("id,title,canonical_name,one_liner,description,hostname,website,url,tags,score,created_at,updated_at,evidence_urls,dedupe_key")\
            .order("updated_at", desc=True)\
            .limit(50)\
            .execute()
        
        all_deals = response.data if hasattr(response, 'data') else []
        
        # 调试模式：记录 deals 读取结果
        if st.session_state.get("debug_mode", False):
            st.session_state["debug_deals_count"] = len(all_deals)
        
        # 获取所有 deal_actions（用于判断哪些已处理）
        actions_response = supabase_client.table("deal_actions")\
            .select("deal_id, action")\
            .execute()
        
        deal_actions_map = {}
        if hasattr(actions_response, 'data'):
            for action in actions_response.data:
                deal_id = action.get('deal_id')
                if deal_id:
                    if deal_id not in deal_actions_map:
                        deal_actions_map[deal_id] = []
                    deal_actions_map[deal_id].append(action.get('action'))
        
        # 应用过滤
        filtered_deals = []
        for deal in all_deals:
            deal_id = deal.get('id') or deal.get('dedupe_key', '')
            
            # 关键字搜索（对 title/canonical_name/one_liner/hostname/website/url 做匹配）
            if search_query:
                search_lower = search_query.lower()
                title = (deal.get('title') or '').lower()
                canonical_name = (deal.get('canonical_name') or '').lower()
                one_liner = (deal.get('one_liner') or deal.get('description') or '').lower()
                hostname = (deal.get('hostname') or '').lower()
                website = (deal.get('website') or '').lower()
                url = (deal.get('url') or '').lower()
                
                # 如果搜索关键词不在任何字段中，跳过
                if (search_lower not in title and 
                    search_lower not in canonical_name and 
                    search_lower not in one_liner and
                    search_lower not in hostname and
                    search_lower not in website and
                    search_lower not in url):
                    continue
            
            # 筛选条件
            if filter_option == "未处理":
                if deal_id in deal_actions_map:
                    continue  # 已处理，跳过
            elif filter_option == "已标记":
                if deal_id not in deal_actions_map:
                    continue  # 未处理，跳过
            
            filtered_deals.append(deal)
        
        # 调试模式：记录过滤后结果
        if st.session_state.get("debug_mode", False):
            st.session_state["debug_filtered_deals_count"] = len(filtered_deals)
        
        # 显示统计
        st.metric("候选项目", len(filtered_deals))
        
        if not filtered_deals:
            st.info("📭 暂无符合条件的项目")
        else:
            # 显示列表
            for deal in filtered_deals:
                # 使用 deal.id 作为唯一标识（不要用 enumerate 的 idx）
                deal_id = deal.get('id') or deal.get('dedupe_key', '')
                
                # 主标题优先级兜底：deal.title ?? deal.canonical_name ?? deal.one_liner ?? deal.hostname ?? deal.website ?? deal.url ?? "(untitled)"
                main_title = (deal.get('title') or 
                             deal.get('canonical_name') or 
                             deal.get('one_liner') or 
                             deal.get('hostname') or 
                             deal.get('website') or 
                             deal.get('url') or 
                             "(untitled)")
                
                canonical_name = deal.get('canonical_name') or deal.get('title', '')
                one_liner = deal.get('one_liner') or deal.get('description', '')
                website = deal.get('website') or deal.get('url', '')
                updated_at = deal.get('updated_at') or deal.get('created_at', '')
                
                # 判断是否已处理
                has_action = deal_id in deal_actions_map
                action_badge = ""
                if has_action:
                    actions = deal_actions_map[deal_id]
                    action_badge = f" | 已标记: {', '.join(actions)}"
                
                # 使用 expander 显示（st.expander 不支持 key 参数）
                with st.expander(f"📌 {main_title}{action_badge}", expanded=False):
                    # 基本信息
                    col1, col2 = st.columns([2, 1])
                    
                    with col1:
                        st.markdown(f"**项目名称**: {canonical_name}")
                        if website:
                            st.markdown(f"**网站**: [{website}]({website})" if website.startswith('http') else f"**网站**: {website}")
                        if one_liner:
                            # 截断显示
                            one_liner_display = one_liner[:200] + "..." if len(one_liner) > 200 else one_liner
                            st.markdown(f"**简介**: {one_liner_display}")
                        if updated_at:
                            try:
                                if "T" in updated_at:
                                    dt = datetime.fromisoformat(updated_at.replace('Z', '+00:00'))
                                    updated_str = dt.strftime("%Y-%m-%d %H:%M")
                                else:
                                    updated_str = updated_at
                            except:
                                updated_str = updated_at
                            st.caption(f"🕒 更新时间: {updated_str}")
                    
                    with col2:
                        # 显示详情
                        evidence_urls = deal.get('evidence_urls', [])
                        if evidence_urls:
                            st.markdown("**证据链接**:")
                            if isinstance(evidence_urls, list):
                                for url in evidence_urls[:3]:  # 只显示前3个
                                    st.markdown(f"- [{url[:50]}...]({url})" if len(url) > 50 else f"- [{url}]({url})")
                            else:
                                st.text(str(evidence_urls)[:100])
                    
                    st.divider()
                    
                    # 三个按钮（并排）
                    col_btn1, col_btn2, col_btn3, col_btn4 = st.columns(4)
                    
                    with col_btn1:
                        if st.button(f"✅ Intro", key=f"intro_{deal_id}", type="primary"):
                            try:
                                action_data = {
                                    "deal_id": deal_id,
                                    "action": "intro",
                                    "notes": "",
                                    "created_at": datetime.utcnow().isoformat()
                                }
                                supabase_client.table("deal_actions").insert(action_data).execute()
                                
                                # 调试模式：记录写入的 payload
                                if st.session_state.get("debug_mode", False):
                                    payload_copy = action_data.copy()
                                    if payload_copy.get("notes") and len(payload_copy["notes"]) > 50:
                                        payload_copy["notes"] = payload_copy["notes"][:50] + "..."
                                    st.session_state["last_action_payload"] = payload_copy
                                    st.sidebar.success(f"✅ {action_data['action']} | deal_id: {deal_id[:8]}...")
                                
                                st.success("✅ 已记录: Intro")
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ 记录失败: {e}")
                    
                    with col_btn2:
                        if st.button(f"👀 Watch", key=f"watch_{deal_id}", type="secondary"):
                            try:
                                action_data = {
                                    "deal_id": deal_id,
                                    "action": "watch",
                                    "notes": "",
                                    "created_at": datetime.utcnow().isoformat()
                                }
                                supabase_client.table("deal_actions").insert(action_data).execute()
                                
                                # 调试模式：记录写入的 payload
                                if st.session_state.get("debug_mode", False):
                                    payload_copy = action_data.copy()
                                    if payload_copy.get("notes") and len(payload_copy["notes"]) > 50:
                                        payload_copy["notes"] = payload_copy["notes"][:50] + "..."
                                    st.session_state["last_action_payload"] = payload_copy
                                    st.sidebar.success(f"✅ {action_data['action']} | deal_id: {deal_id[:8]}...")
                                
                                st.success("✅ 已记录: Watch")
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ 记录失败: {e}")
                    
                    with col_btn3:
                        if st.button(f"❌ Pass", key=f"pass_{deal_id}", type="secondary"):
                            try:
                                action_data = {
                                    "deal_id": deal_id,
                                    "action": "pass",
                                    "notes": "",
                                    "created_at": datetime.utcnow().isoformat()
                                }
                                supabase_client.table("deal_actions").insert(action_data).execute()
                                
                                # 调试模式：记录写入的 payload
                                if st.session_state.get("debug_mode", False):
                                    payload_copy = action_data.copy()
                                    if payload_copy.get("notes") and len(payload_copy["notes"]) > 50:
                                        payload_copy["notes"] = payload_copy["notes"][:50] + "..."
                                    st.session_state["last_action_payload"] = payload_copy
                                    st.sidebar.success(f"✅ {action_data['action']} | deal_id: {deal_id[:8]}...")
                                
                                st.success("✅ 已记录: Pass")
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ 记录失败: {e}")
                    
                    with col_btn4:
                        # 显示已有的操作记录
                        if has_action:
                            actions = deal_actions_map[deal_id]
                            st.caption(f"已操作: {', '.join(actions)}")
    
    except Exception as e:
        st.error(f"❌ 加载数据失败: {e}")
        st.exception(e)

elif page == "📄 周报":
    # ========== 周报页面 ==========
    st.title("📄 周报")
    st.caption("查看最新生成的雷达周报")
    
    if not HAS_SUPABASE:
        st.error("❌ Supabase 客户端未安装。请运行: `pip install supabase`")
        st.stop()
    
    # 获取 Supabase 客户端（前端只使用 ANON_KEY）
    supabase_client = get_supabase_client(use_service_role=False)
    
    if not supabase_client:
        st.error("❌ Supabase 连接失败。请检查环境变量：")
        st.code("""
SUPABASE_URL=your_supabase_url
SUPABASE_ANON_KEY=your_supabase_anon_key
        """)
        st.info("💡 请在 `.env` 文件中设置这些环境变量，或使用系统环境变量。")
        st.stop()
    
    try:
        # 获取最新一条周报（按 week_start 倒序）
        response = supabase_client.table("weekly_reports")\
            .select("*")\
            .order("week_start", desc=True)\
            .limit(1)\
            .execute()
        
        reports = response.data if hasattr(response, 'data') else []
        
        # 调试模式：记录周报读取结果
        if st.session_state.get("debug_mode", False):
            if not reports:
                st.session_state["debug_weekly_report"] = "no reports"
            else:
                report = reports[0]
                week_start = report.get('week_start', '')
                # 优先读取 markdown 字段，如果不存在则回退到 content（兼容旧数据）
                content = report.get('markdown', '') or report.get('content', '')
                st.session_state["debug_weekly_report"] = {
                    "week_start": week_start,
                    "markdown_length": len(content) if content else 0
                }
        
        if not reports:
            st.info("📭 暂无周报。雷达抓取任务运行后，周报会自动生成。")
            st.info("💡 周报会在每周一和周三自动生成（通过 GitHub Actions）。")
        else:
            report = reports[0]
            week_start = report.get('week_start', '')
            # 优先读取 markdown 字段，如果不存在则回退到 content（兼容旧数据）
            content = report.get('markdown', '') or report.get('content', '')
            created_at = report.get('created_at', '')
            
            # 解析 week_start 日期
            try:
                if "T" in week_start:
                    week_dt = datetime.fromisoformat(week_start.replace('Z', '+00:00'))
                    week_start_display = week_dt.strftime("%Y-%m-%d")
                    week_end_display = (week_dt + timedelta(days=6)).strftime("%Y-%m-%d")
                else:
                    week_start_display = week_start
                    week_end_display = "未知"
            except:
                week_start_display = week_start
                week_end_display = "未知"
            
            # 显示周报信息
            col1, col2 = st.columns([2, 1])
            with col1:
                st.markdown(f"**周报期间**: {week_start_display} 至 {week_end_display}")
            
            with col2:
                if created_at:
                    try:
                        if "T" in created_at:
                            created_dt = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                            created_display = created_dt.strftime("%Y-%m-%d %H:%M")
                        else:
                            created_display = created_at
                    except:
                        created_display = created_at
                    st.caption(f"生成时间: {created_display}")
            
            st.divider()
            
            # 下载按钮
            download_filename = f"weekly_report_{week_start_display.replace('-', '')}.md"
            st.download_button(
                label="⬇️ 下载周报 (Markdown)",
                data=content,
                file_name=download_filename,
                mime="text/markdown",
                key="download_weekly_report"
            )
            
            st.divider()
            
            # 渲染 Markdown 内容
            st.markdown(content)
    
    except Exception as e:
        st.error(f"❌ 加载周报失败: {e}")
        st.exception(e)
